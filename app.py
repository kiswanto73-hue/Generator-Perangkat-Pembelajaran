import streamlit as st
from groq import Groq
import re
import requests
import pandas as pd
from datetime import datetime, date, timedelta
from io import BytesIO

from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, Image as RLImage
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import qrcode

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from supabase import create_client
except Exception:
    create_client = None

# ============================================================
# KONFIGURASI HALAMAN
# ============================================================
st.set_page_config(page_title="Generator Perangkat Ajar KKG", layout="wide", page_icon="📚")

MODEL_GROQ = "openai/gpt-oss-120b"

COLOR_PRIMARY = colors.HexColor("#1E4D8C")
COLOR_ACCENT = colors.HexColor("#F2A900")
COLOR_LIGHT = colors.HexColor("#EEF3FA")
DOCX_PRIMARY_RGB = RGBColor(0x1E, 0x4D, 0x8C)
DOCX_GREY_RGB = RGBColor(0x55, 0x55, 0x55)

# ============================================================
# TEMA VISUAL - CSS ELEGAN (diterapkan ke semua halaman/tab)
# ============================================================
st.markdown("""
<style>
    .stApp { background: linear-gradient(180deg, #F7F9FC 0%, #FFFFFF 100%); }
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1E4D8C 0%, #163a6b 100%);
    }
    section[data-testid="stSidebar"] * { color: #F2F5FA !important; }
    section[data-testid="stSidebar"] input, section[data-testid="stSidebar"] textarea {
        color: #16324F !important;
    }
    h1, h2, h3 { color: #1E4D8C; font-family: "Source Sans Pro", sans-serif; }
    div.stButton > button {
        background: linear-gradient(135deg, #1E4D8C 0%, #2E6BB8 100%);
        color: white; border: none; border-radius: 8px; font-weight: 600;
        padding: 0.5rem 1.1rem; transition: 0.15s ease-in-out;
    }
    div.stButton > button:hover { transform: translateY(-1px); box-shadow: 0 4px 10px rgba(30,77,140,0.35); }
    div[data-testid="stExpander"], div[data-baseweb="tab-panel"] {
        background: #FFFFFF; border-radius: 12px;
    }
    .stTabs [data-baseweb="tab-list"] { gap: 6px; }
    .stTabs [data-baseweb="tab"] {
        background-color: #EEF3FA; border-radius: 8px 8px 0 0; padding: 8px 16px; font-weight: 600;
    }
    .stTabs [aria-selected="true"] { background-color: #1E4D8C !important; color: white !important; }
    div[data-testid="stMetric"], .kkg-card {
        background: #FFFFFF; border: 1px solid #E3E9F2; border-radius: 12px;
        padding: 14px 18px; box-shadow: 0 2px 8px rgba(30,77,140,0.06);
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# PEDOMAN PEMBELAJARAN MENDALAM (DEEP LEARNING)
# Permendikdasmen No. 13 Tahun 2025 - penyempurnaan Permendikbudristek No. 12/2024
# ============================================================
DEEP_LEARNING_GUIDE = """
PEDOMAN WAJIB - PARADIGMA PEMBELAJARAN MENDALAM (DEEP LEARNING)
Sesuai Permendikdasmen No. 13 Tahun 2025, seluruh rancangan pembelajaran harus
menghadirkan tiga prinsip berikut secara EKSPLISIT (beri label di setiap langkah kegiatan inti):

1. MINDFUL (Berkesadaran): peserta didik diajak fokus penuh, sadar akan tujuan
   belajarnya, serta diberi ruang refleksi diri di awal/akhir kegiatan.
2. MEANINGFUL (Bermakna): materi dikaitkan secara kontekstual dengan pengalaman
   nyata dan kehidupan sehari-hari murid, bukan sekadar hafalan.
3. JOYFUL (Menyenangkan): suasana belajar positif, melibatkan unsur kolaborasi,
   permainan, eksplorasi, atau apresiasi karya, sehingga murid termotivasi.

Setiap langkah pada bagian Kegiatan Inti WAJIB diberi penanda, contoh:
"Guru mengajak murid duduk tenang sejenak dan menyebutkan tujuan belajar hari ini. (Mindful)"

ATURAN FORMAT KETAT (WAJIB DIPATUHI):
Dokumen ini akan dicetak ke PDF/Word, BUKAN ditampilkan sebagai halaman web. Karena itu:
- JANGAN PERNAH menulis tag HTML apa pun, termasuk <textarea>, <img>, <br>, <div>, <input>,
  <table>, atau tag lainnya. Tag HTML akan muncul sebagai teks aneh/rusak di dokumen cetak.
- Untuk kolom isian jawaban, gunakan garis titik-titik biasa, contoh: Jawaban: ______________
- Untuk kotak kosong isian, gunakan format: [ ................................. ]
- Untuk ilustrasi/gambar yang perlu ditempel guru, TULISKAN keterangannya dalam kalimat biasa
  di dalam tanda kurung siku, contoh: [Petunjuk ilustrasi: gambar dua ekor kucing bermain bola].
  JANGAN PERNAH menyertakan tag <img>, URL gambar, atau tautan gambar apa pun (termasuk dari
  placeholder.com atau sejenisnya) — gambar sungguhan tidak bisa dihasilkan lewat teks.
- Gunakan HANYA format Markdown biasa (heading #, bold **, tabel markdown, list -), tanpa HTML.
"""

SUBJECT_OPTIONS = ["PAI & BP", "Bahasa Indonesia", "Matematika", "IPAS", "PJOK",
                    "Pendidikan Pancasila", "Seni Rupa", "Bahasa Inggris", "Lainnya"]

# Input Kelas 1-6 (setiap kelas otomatis dipasangkan dengan Fase Kurikulum Merdeka)
KELAS_FASE_MAP = {
    "Kelas 1": "Fase A", "Kelas 2": "Fase A",
    "Kelas 3": "Fase B", "Kelas 4": "Fase B",
    "Kelas 5": "Fase C", "Kelas 6": "Fase C",
}
KELAS_OPTIONS = list(KELAS_FASE_MAP.keys())


def label_kelas_fase(kelas: str) -> str:
    return f"{kelas} ({KELAS_FASE_MAP.get(kelas, '')})"


# ============================================================
# STATE AWAL
# ============================================================
defaults = {
    "sekolah": "SDN 165/V Tanjung Jabung Barat",
    "penyusun": "",
    "guru_nip": "",
    "kepsek_nama": "",
    "kepsek_nip": "",
    "tahun_ajaran": "2025/2026",
    "hasil_tp_atp": "", "meta_tp_atp": {},
    "hasil_modul": "", "meta_modul": {},
    "hasil_lkpd": "", "meta_lkpd": {}, "lkpd_gambar_bytes": None,
    "hasil_prota_promes": "", "meta_prota_promes": {}, "judul_prota_promes": "PROGRAM TAHUNAN (PROTA)",
    "hasil_minggu_efektif": "", "meta_minggu_efektif": {}, "tabel_minggu_efektif": None,
    "jurnal_rows": [],
    "siswa_rows": [],
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v
if "user" not in st.session_state:
    st.session_state.user = None
if "data_loaded" not in st.session_state:
    st.session_state.data_loaded = False


# ============================================================
# AUTENTIKASI & PENYIMPANAN PERMANEN (SUPABASE) - OPSIONAL
# Aktif otomatis jika SUPABASE_URL & SUPABASE_KEY diisi di .streamlit/secrets.toml.
# Jika tidak diisi, aplikasi tetap berjalan seperti biasa (data hanya tersimpan
# selama sesi browser terbuka, tanpa login) - lihat README_DEPLOY.md.
# ============================================================
SUPABASE_URL = st.secrets.get("SUPABASE_URL", "") if hasattr(st, "secrets") else ""
SUPABASE_KEY = st.secrets.get("SUPABASE_KEY", "") if hasattr(st, "secrets") else ""
AUTH_AKTIF = bool(SUPABASE_URL and SUPABASE_KEY and create_client is not None)

# Kunci khusus admin (service_role) & kata sandi panel admin - HANYA diisi oleh pemilik
# aplikasi di Secrets, TIDAK PERNAH dikirim ke browser guru. Dipakai untuk membuat &
# memantau kode lisensi. Jika kosong, panel admin otomatis nonaktif.
SUPABASE_SERVICE_KEY = st.secrets.get("SUPABASE_SERVICE_KEY", "") if hasattr(st, "secrets") else ""
ADMIN_PASSWORD = st.secrets.get("ADMIN_PASSWORD", "") if hasattr(st, "secrets") else ""
ADMIN_AKTIF = bool(SUPABASE_URL and SUPABASE_SERVICE_KEY and ADMIN_PASSWORD and create_client is not None)


@st.cache_resource
def get_supabase():
    return create_client(SUPABASE_URL, SUPABASE_KEY)


@st.cache_resource
def get_supabase_admin():
    """Client khusus admin, pakai service_role key -> bisa lewati RLS sepenuhnya.
    Hanya dipanggil dari panel admin yang sudah dilindungi kata sandi."""
    return create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)


def generate_kode_unik() -> str:
    import secrets as _secrets, string as _string
    abjad = _string.ascii_uppercase + _string.digits
    blok = lambda n: "".join(_secrets.choice(abjad) for _ in range(n))
    return f"KKG-{blok(4)}-{blok(4)}"


def rpc_reserve_kode(kode: str) -> bool:
    """Kunci kode secara atomik lewat anon key (aman dari race condition)."""
    sb = get_supabase()
    res = sb.rpc("reserve_kode_lisensi", {"p_kode": kode}).execute()
    return bool(res.data)


def rpc_batalkan_kode(kode: str):
    try:
        get_supabase().rpc("batalkan_kode_lisensi", {"p_kode": kode}).execute()
    except Exception:
        pass


def rpc_tautkan_kode(kode: str, user_id: str, email: str):
    try:
        get_supabase().rpc("tautkan_kode_lisensi", {"p_kode": kode, "p_user_id": user_id, "p_email": email}).execute()
    except Exception:
        pass


# Notifikasi WhatsApp otomatis ke admin saat ada pendaftaran guru baru (via Fonnte).
# Isi TOKEN & NOMOR di Secrets (bukan di kode), agar tidak ikut ter-upload ke GitHub publik.
TOKEN_FONNTE = st.secrets.get("eRiRwGfU96GqAyaMXM7Q", "") if hasattr(st, "secrets") else ""
NOMOR_WA_ADMIN = st.secrets.get("6282177723494", "") if hasattr(st, "secrets") else ""
WA_NOTIF_AKTIF = bool(TOKEN_FONNTE and NOMOR_WA_ADMIN)


def kirim_wa_fonnte(nomor: str, pesan: str, token: str) -> bool:
    """Kirim pesan WA lewat Fonnte. Dibungkus try/except agar kegagalan kirim WA
    (kuota habis, token salah, dll) TIDAK sampai menggagalkan pendaftaran guru."""
    try:
        resp = requests.post(
            "https://api.fonnte.com/send",
            headers={"Authorization": token},
            data={"target": nomor, "message": pesan},
            timeout=10,
        )
        return resp.status_code == 200
    except Exception:
        return False


def notifikasi_wa_pendaftaran_baru(nama: str, email: str, kode: str):
    if not WA_NOTIF_AKTIF:
        return
    pesan_admin = (
        "*PENDAFTARAN AKUN BARU!*\n"
        "Ada guru yang baru saja mendaftar di aplikasi:\n"
        f"👤 *Nama:* {nama}\n"
        f"📧 *Email:* {email}\n"
        f"🔑 *Kode Aktivasi yang Digunakan:* {kode}\n"
        "Silakan catat data pendaftar ini."
    )
    kirim_wa_fonnte(NOMOR_WA_ADMIN, pesan_admin, TOKEN_FONNTE)


# --- Konfigurasi pembayaran & permintaan kode mandiri lewat poster promosi ---
DANA_NOMOR = st.secrets.get("DANA_NOMOR", "") if hasattr(st, "secrets") else ""
DANA_NAMA = st.secrets.get("DANA_NAMA", "") if hasattr(st, "secrets") else ""
HARGA_KODE = st.secrets.get("HARGA_KODE", "") if hasattr(st, "secrets") else ""
WA_ADMIN_TAMPIL = st.secrets.get("NOMOR_WA_ADMIN", "") if hasattr(st, "secrets") else ""
POSTER_AKTIF = bool(DANA_NOMOR and DANA_NAMA and HARGA_KODE and WA_ADMIN_TAMPIL)


def buat_dan_kirim_permintaan_kode(nama_pemohon: str, sekolah_pemohon: str, wa_pemohon: str):
    """Dipanggil dari halaman poster publik (belum login). Generate 1 kode baru & simpan ke
    database lewat Service Key (aman, server-side), lalu kirim ke WA ADMIN (bukan ke pemohon
    langsung) supaya Admin bisa verifikasi bukti transfer dulu sebelum meneruskan kodenya."""
    if not ADMIN_AKTIF:
        return False, "Sistem otomatis belum aktif (Admin belum atur Secrets)."
    kode_baru = generate_kode_unik()
    catatan = f"Permintaan mandiri via poster - {nama_pemohon} ({sekolah_pemohon}) - WA {wa_pemohon}"
    try:
        get_supabase_admin().table("kode_lisensi").insert(
            {"kode": kode_baru, "catatan": catatan}).execute()
    except Exception as e:
        return False, f"Gagal menyimpan kode baru: {e}"

    pesan = (
        "*PERMINTAAN KODE AKTIVASI BARU!*\n"
        "Ada guru yang mengaku sudah transfer mahar dan minta kode aktivasi:\n"
        f"👤 *Nama:* {nama_pemohon}\n"
        f"🏫 *Sekolah:* {sekolah_pemohon}\n"
        f"📱 *No. WA Pemohon:* {wa_pemohon}\n"
        f"🔑 *Kode Aktivasi (BARU, siap dipakai):* {kode_baru}\n\n"
        "⚠️ Cek dulu bukti transfer dari nomor WA di atas sebelum kode ini diteruskan ke yang bersangkutan."
    )
    kirim_wa_fonnte(NOMOR_WA_ADMIN, pesan, TOKEN_FONNTE)
    return True, kode_baru


def sb_daftar(email: str, password: str, nama: str):
    sb = get_supabase()
    res = sb.auth.sign_up({"email": email, "password": password,
                            "options": {"data": {"nama_lengkap": nama}}})
    return res


def sb_masuk(email: str, password: str):
    sb = get_supabase()
    res = sb.auth.sign_in_with_password({"email": email, "password": password})
    return res


def sb_keluar():
    try:
        get_supabase().auth.sign_out()
    except Exception:
        pass
    st.session_state.user = None
    st.session_state.data_loaded = False
    st.session_state.jurnal_rows = []
    st.session_state.siswa_rows = []


def db_muat_data(user_id: str):
    """Muat jurnal & data siswa milik guru yang sedang login dari Supabase."""
    sb = get_supabase()
    try:
        j = sb.table("jurnal_mengajar").select("*").eq("user_id", user_id).order("id").execute()
        st.session_state.jurnal_rows = [
            {"id": r["id"], "Tanggal": r["tanggal"], "Kelas": r["kelas"], "Mapel": r["mapel"],
             "JP": r["jp"], "Materi": r["materi"], "Kegiatan": r["kegiatan"],
             "Hadir": r["hadir"], "Catatan": r["catatan"]}
            for r in (j.data or [])
        ]
    except Exception as e:
        st.warning(f"⚠️ Gagal memuat data jurnal dari database: {e}")
    try:
        s = sb.table("data_siswa").select("*").eq("user_id", user_id).order("id").execute()
        st.session_state.siswa_rows = [
            {"id": r["id"], "Nama": r["nama"], "L/P": r["jk"], "NISN": r["nisn"]}
            for r in (s.data or [])
        ]
    except Exception as e:
        st.warning(f"⚠️ Gagal memuat data siswa dari database: {e}")


def db_tambah_jurnal(user_id: str, row: dict):
    sb = get_supabase()
    payload = {"user_id": user_id, "tanggal": row["Tanggal"], "kelas": row["Kelas"],
               "mapel": row["Mapel"], "jp": row["JP"], "materi": row["Materi"],
               "kegiatan": row["Kegiatan"], "hadir": row["Hadir"], "catatan": row["Catatan"]}
    try:
        res = sb.table("jurnal_mengajar").insert(payload).execute()
        if res.data:
            row["id"] = res.data[0]["id"]
    except Exception as e:
        st.warning(f"⚠️ Gagal menyimpan jurnal ke database: {e}")
    return row


def db_hapus_jurnal(row_id):
    if row_id is None:
        return
    try:
        get_supabase().table("jurnal_mengajar").delete().eq("id", row_id).execute()
    except Exception as e:
        st.warning(f"⚠️ Gagal menghapus data di database: {e}")


def db_kosongkan_jurnal(user_id: str):
    try:
        get_supabase().table("jurnal_mengajar").delete().eq("user_id", user_id).execute()
    except Exception as e:
        st.warning(f"⚠️ Gagal mengosongkan data di database: {e}")


def db_tambah_siswa(user_id: str, row: dict):
    sb = get_supabase()
    payload = {"user_id": user_id, "nama": row["Nama"], "jk": row["L/P"], "nisn": row.get("NISN", "")}
    try:
        res = sb.table("data_siswa").insert(payload).execute()
        if res.data:
            row["id"] = res.data[0]["id"]
    except Exception as e:
        st.warning(f"⚠️ Gagal menyimpan data siswa ke database: {e}")
    return row


def db_hapus_siswa(row_id):
    if row_id is None:
        return
    try:
        get_supabase().table("data_siswa").delete().eq("id", row_id).execute()
    except Exception as e:
        st.warning(f"⚠️ Gagal menghapus data di database: {e}")


def db_kosongkan_siswa(user_id: str):
    try:
        get_supabase().table("data_siswa").delete().eq("user_id", user_id).execute()
    except Exception as e:
        st.warning(f"⚠️ Gagal mengosongkan data di database: {e}")


def db_catat_absensi_barcode(user_id: str, nisn: str, nama: str, kelas: str,
                               tanggal_str: str, status: str = "Hadir"):
    """Catat/perbarui kehadiran satu siswa untuk satu tanggal (upsert berdasarkan
    kombinasi guru+NISN+tanggal, jadi aman di-scan ulang tanpa duplikat)."""
    sb = get_supabase()
    payload = {"user_id": user_id, "nisn": nisn, "nama": nama, "kelas": kelas,
               "tanggal": tanggal_str, "status": status}
    try:
        sb.table("absensi_barcode").upsert(payload, on_conflict="user_id,nisn,tanggal").execute()
        return True, None
    except Exception as e:
        return False, str(e)


def db_ambil_absensi_tanggal(user_id: str, tanggal_str: str):
    try:
        res = (get_supabase().table("absensi_barcode").select("*")
               .eq("user_id", user_id).eq("tanggal", tanggal_str)
               .order("waktu", desc=True).execute())
        return res.data or []
    except Exception:
        return []


def db_ambil_absensi_rentang(user_id: str, tgl_awal: str, tgl_akhir: str):
    try:
        res = (get_supabase().table("absensi_barcode").select("*")
               .eq("user_id", user_id).gte("tanggal", tgl_awal).lte("tanggal", tgl_akhir)
               .order("tanggal").execute())
        return res.data or []
    except Exception:
        return []


def buat_gambar_qr_nisn(nisn: str) -> BytesIO:
    """Hasilkan gambar QR Code dari NISN, dalam memori (tidak disimpan ke disk/DB)."""
    buf = BytesIO()
    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_M,
                        box_size=10, border=2)
    qr.add_data(nisn)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def tambah_baris_jurnal(row: dict):
    if AUTH_AKTIF and st.session_state.user:
        row = db_tambah_jurnal(st.session_state.user["id"], row)
    st.session_state.jurnal_rows.append(row)


def hapus_jurnal_terakhir():
    if st.session_state.jurnal_rows:
        row = st.session_state.jurnal_rows.pop()
        if AUTH_AKTIF and st.session_state.user:
            db_hapus_jurnal(row.get("id"))


def kosongkan_jurnal():
    if AUTH_AKTIF and st.session_state.user:
        db_kosongkan_jurnal(st.session_state.user["id"])
    st.session_state.jurnal_rows = []


def tambah_baris_siswa(row: dict):
    if AUTH_AKTIF and st.session_state.user:
        row = db_tambah_siswa(st.session_state.user["id"], row)
    st.session_state.siswa_rows.append(row)


def hapus_siswa_terakhir():
    if st.session_state.siswa_rows:
        row = st.session_state.siswa_rows.pop()
        if AUTH_AKTIF and st.session_state.user:
            db_hapus_siswa(row.get("id"))


def kosongkan_siswa():
    if AUTH_AKTIF and st.session_state.user:
        db_kosongkan_siswa(st.session_state.user["id"])
    st.session_state.siswa_rows = []


def tampilkan_gerbang_login():
    """Tampilkan formulir Masuk/Daftar. Menghentikan eksekusi (st.stop()) sampai berhasil login."""
    st.title("📚 Generator Perangkat Ajar Kurikulum Merdeka")
    st.caption("Silakan masuk atau daftar akun guru untuk mulai menggunakan aplikasi. "
               "Data jurnal & absen Anda akan tersimpan permanen dan hanya bisa diakses oleh akun Anda sendiri.")
    tab_masuk, tab_daftar, tab_beli = st.tabs(
        ["🔑 Masuk", "📝 Daftar Akun Baru", "🛒 Belum Punya Kode?"])

    with tab_masuk:
        with st.form("form_masuk"):
            email_masuk = st.text_input("Email", key="email_masuk")
            pw_masuk = st.text_input("Kata Sandi", type="password", key="pw_masuk")
            submit_masuk = st.form_submit_button("🔑 Masuk", use_container_width=True)
        if submit_masuk:
            if not email_masuk or not pw_masuk:
                st.error("⚠️ Mohon isi email dan kata sandi.")
            else:
                try:
                    res = sb_masuk(email_masuk, pw_masuk)
                    if res.user:
                        st.session_state.user = {
                            "id": res.user.id, "email": res.user.email,
                            "access_token": res.session.access_token if res.session else "",
                        }
                        st.success("✅ Berhasil masuk!")
                        st.rerun()
                    else:
                        st.error("❌ Email atau kata sandi salah.")
                except Exception as e:
                    st.error(f"❌ Gagal masuk: {e}")

    with tab_daftar:
        st.caption("Butuh **kode aktivasi** dari pemilik aplikasi untuk mendaftar. "
                    "Satu kode hanya berlaku untuk satu akun guru, berlaku selamanya.")
        with st.form("form_daftar"):
            nama_daftar = st.text_input("Nama Lengkap", key="nama_daftar")
            email_daftar = st.text_input("Email", key="email_daftar")
            pw_daftar = st.text_input("Kata Sandi (minimal 6 karakter)", type="password", key="pw_daftar")
            kode_daftar = st.text_input("Kode Aktivasi", key="kode_daftar",
                                          placeholder="contoh: KKG-A1B2-C3D4").strip().upper()
            submit_daftar = st.form_submit_button("📝 Daftar", use_container_width=True)
        if submit_daftar:
            if not nama_daftar or not email_daftar or not pw_daftar or not kode_daftar:
                st.error("⚠️ Mohon lengkapi semua isian, termasuk kode aktivasi.")
            elif len(pw_daftar) < 6:
                st.error("⚠️ Kata sandi minimal 6 karakter.")
            else:
                try:
                    kode_valid = rpc_reserve_kode(kode_daftar)
                except Exception as e:
                    kode_valid = False
                    st.error(f"❌ Gagal memeriksa kode aktivasi: {e}")
                if not kode_valid:
                    st.error("❌ Kode aktivasi tidak ditemukan atau sudah pernah digunakan. "
                              "Hubungi pemilik aplikasi untuk mendapatkan kode.")
                else:
                    try:
                        res = sb_daftar(email_daftar, pw_daftar, nama_daftar)
                        if res.user:
                            rpc_tautkan_kode(kode_daftar, res.user.id, email_daftar)
                            notifikasi_wa_pendaftaran_baru(nama_daftar, email_daftar, kode_daftar)
                            st.success("✅ Pendaftaran berhasil! Jika verifikasi email diaktifkan, "
                                        "silakan cek email Anda terlebih dahulu, lalu masuk lewat tab 'Masuk'.")
                        else:
                            rpc_batalkan_kode(kode_daftar)
                            st.error("❌ Pendaftaran gagal, coba lagi.")
                    except Exception as e:
                        rpc_batalkan_kode(kode_daftar)
                        st.error(f"❌ Gagal mendaftar: {e}")

    with tab_beli:
        if not POSTER_AKTIF:
            st.warning("⚠️ Halaman ini belum diatur oleh Admin (nomor DANA, nama pemilik "
                       "rekening, harga, dan nomor WA admin belum lengkap di Secrets).")
        else:
            wa_link_pesan = (
                "Halo Admin, saya guru dan ingin mendapatkan kode aktivasi "
                "Generator Perangkat Ajar KKG. Ini bukti transfer saya:"
            )
            wa_link = f"https://wa.me/{WA_ADMIN_TAMPIL}?text={requests.utils.quote(wa_link_pesan)}"
            st.markdown(f"""
<div style="background:linear-gradient(135deg,#1E4D8C,#3E7CC7);border-radius:16px;
            padding:28px 24px;color:white;text-align:center;margin-bottom:18px;">
  <div style="font-size:14px;letter-spacing:2px;opacity:.85;">GENERATOR PERANGKAT AJAR • KURIKULUM MERDEKA</div>
  <div style="font-size:28px;font-weight:800;margin:8px 0 4px;">✨ Akses Selamanya, Sekali Bayar ✨</div>
  <div style="font-size:15px;opacity:.9;">TP • ATP • Modul Ajar • LKPD • Absen Kartu QR • Jurnal Digital</div>
  <div style="font-size:38px;font-weight:900;margin:18px 0 4px;">{HARGA_KODE}</div>
  <div style="font-size:13px;opacity:.85;">sekali bayar, akun aktif SELAMANYA — tanpa langganan bulanan</div>
</div>
""", unsafe_allow_html=True)

            colp1, colp2 = st.columns(2)
            with colp1:
                st.markdown("#### 1️⃣ Transfer ke DANA")
                st.markdown(f"""
<div style="border:2px dashed #1E4D8C;border-radius:12px;padding:16px;text-align:center;">
  <div style="font-size:13px;color:#555;">Nomor DANA</div>
  <div style="font-size:22px;font-weight:800;color:#1E4D8C;">{DANA_NOMOR}</div>
  <div style="font-size:14px;color:#333;margin-top:4px;">a.n. {DANA_NAMA}</div>
</div>
""", unsafe_allow_html=True)
            with colp2:
                st.markdown("#### 2️⃣ Kirim Bukti ke WA Admin")
                st.link_button("💬 Chat Admin di WhatsApp", wa_link, use_container_width=True)
                st.caption(f"Atau simpan nomor: {WA_ADMIN_TAMPIL}")

            st.divider()
            st.markdown("#### 3️⃣ Sudah Transfer? Kirim Permintaan Kode di Sini")
            st.caption("Setelah admin cek bukti transfer dari WA Anda, kode aktivasi akan "
                       "dikirim balik ke nomor WA yang Anda isi di bawah.")
            with st.form("form_minta_kode"):
                nama_pemohon = st.text_input("Nama Lengkap", key="nama_pemohon")
                sekolah_pemohon = st.text_input("Asal Sekolah", key="sekolah_pemohon")
                wa_pemohon = st.text_input("No. WA Aktif (yang dipakai transfer/chat admin)",
                                            key="wa_pemohon", placeholder="08xxxxxxxxxx")
                kirim = st.form_submit_button("📩 Sudah Transfer, Kirim Permintaan Kode",
                                               use_container_width=True)
            if kirim:
                if not nama_pemohon or not sekolah_pemohon or not wa_pemohon:
                    st.error("⚠️ Mohon lengkapi semua isian.")
                else:
                    ok, hasil = buat_dan_kirim_permintaan_kode(nama_pemohon, sekolah_pemohon, wa_pemohon)
                    if ok:
                        st.success("✅ Permintaan terkirim! Admin akan memverifikasi pembayaran Anda "
                                   "lalu mengirim kode aktivasi ke WhatsApp yang Anda isi.")
                    else:
                        st.error(f"❌ {hasil}")
    st.stop()


def tampilkan_panel_admin():
    """Panel khusus pemilik aplikasi: generate & pantau kode lisensi.
    Diakses lewat URL ...streamlit.app/?admin=1 - tidak muncul di navigasi guru."""
    st.title("🛡️ Panel Admin — Kode Lisensi")
    if not ADMIN_AKTIF:
        st.error("Panel admin belum aktif. Lengkapi `SUPABASE_SERVICE_KEY` dan `ADMIN_PASSWORD` "
                 "di Secrets terlebih dahulu (lihat README_DEPLOY.md).")
        st.stop()

    if not st.session_state.get("admin_ok"):
        with st.form("form_admin_login"):
            pw = st.text_input("Kata Sandi Admin", type="password")
            ok = st.form_submit_button("Masuk sebagai Admin")
        if ok:
            if pw == ADMIN_PASSWORD:
                st.session_state.admin_ok = True
                st.rerun()
            else:
                st.error("❌ Kata sandi salah.")
        st.stop()

    sb_admin = get_supabase_admin()

    st.markdown("### ➕ Buat Kode Lisensi Baru")
    with st.form("form_generate_kode"):
        jumlah = st.number_input("Jumlah kode dibuat sekaligus", 1, 100, 1)
        catatan = st.text_input("Catatan (misal: nama sekolah/pembeli)", "")
        buat = st.form_submit_button("🎟️ Generate Kode")
    if buat:
        kode_baru = []
        for _ in range(int(jumlah)):
            k = generate_kode_unik()
            try:
                sb_admin.table("kode_lisensi").insert({"kode": k, "catatan": catatan}).execute()
                kode_baru.append(k)
            except Exception as e:
                st.error(f"❌ Gagal menyimpan kode {k}: {e}")
        if kode_baru:
            st.success(f"✅ {len(kode_baru)} kode berhasil dibuat.")
            st.code("\n".join(kode_baru), language=None)

    st.divider()
    st.markdown("### 📋 Daftar Kode Lisensi")
    filter_status = st.selectbox("Filter status", ["Semua", "belum_terpakai", "terpakai"])
    try:
        q = sb_admin.table("kode_lisensi").select("*").order("dibuat_tanggal", desc=True)
        if filter_status != "Semua":
            q = q.eq("status", filter_status)
        data = q.execute().data or []
    except Exception as e:
        st.error(f"❌ Gagal memuat daftar kode: {e}")
        data = []

    if data:
        df = pd.DataFrame(data)[["kode", "status", "dipakai_oleh_email", "catatan",
                                  "dibuat_tanggal", "digunakan_tanggal"]]
        st.dataframe(df, use_container_width=True, hide_index=True)
        c1, c2 = st.columns(2)
        c1.metric("Total kode", len(data))
        c2.metric("Sudah terpakai", sum(1 for d in data if d["status"] == "terpakai"))
        st.download_button("⬇️ Unduh sebagai CSV", df.to_csv(index=False).encode("utf-8"),
                            "kode_lisensi.csv", "text/csv")
    else:
        st.info("Belum ada kode dibuat.")

    if st.button("🚪 Keluar dari Panel Admin"):
        st.session_state.admin_ok = False
        st.rerun()
    st.stop()


if "admin" in st.query_params:
    tampilkan_panel_admin()

if AUTH_AKTIF and st.session_state.user is None:
    tampilkan_gerbang_login()

if AUTH_AKTIF and st.session_state.user is not None and not st.session_state.data_loaded:
    db_muat_data(st.session_state.user["id"])
    st.session_state.data_loaded = True


# ============================================================
# FUNGSI EKSTRAKSI FILE REFERENSI (Buku Guru PDF/DOCX/TXT/XLSX/CSV)
# ============================================================
def ekstrak_teks_referensi(uploaded_file, max_chars: int = 8000) -> tuple:
    """Ambil cuplikan teks dari file referensi (PDF/DOCX/TXT/XLSX/CSV) untuk dijadikan
    konteks tambahan bagi AI. Mengembalikan (teks, pesan_status)."""
    if uploaded_file is None:
        return "", ""
    nama = uploaded_file.name.lower()
    try:
        data = uploaded_file.getvalue()  # ambil bytes secara utuh agar tidak bentrok posisi pointer
    except Exception as e:
        return "", f"❌ Gagal membaca berkas: {e}"

    try:
        if nama.endswith(".pdf"):
            if PdfReader is None:
                return "", "❌ Modul pembaca PDF tidak tersedia."
            reader = PdfReader(BytesIO(data))
            n_halaman = len(reader.pages)
            potongan = []
            for page in reader.pages[:30]:
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                if t.strip():
                    potongan.append(t)
            teks = "\n".join(potongan).strip()
            if not teks:
                return "", (f"⚠️ Berkas PDF ({n_halaman} halaman) terbaca tapi TIDAK ADA teks yang bisa "
                             f"diekstrak — kemungkinan ini adalah hasil SCAN/FOTO (gambar), bukan PDF teks. "
                             f"Silakan gunakan PDF asli (bukan hasil scan) atau salin manual isinya ke kolom teks.")
            return teks[:max_chars], f"✅ Berhasil membaca {len(teks)} karakter dari {n_halaman} halaman PDF."

        if nama.endswith(".docx"):
            doc = Document(BytesIO(data))
            bagian = [p.text for p in doc.paragraphs if p.text.strip()]
            for tabel in doc.tables:
                for row in tabel.rows:
                    bagian.append(" | ".join(c.text for c in row.cells))
            teks = "\n".join(bagian).strip()
            if not teks:
                return "", "⚠️ Berkas DOCX terbaca tapi tidak ada teks di dalamnya (kemungkinan hanya berisi gambar)."
            return teks[:max_chars], f"✅ Berhasil membaca {len(teks)} karakter dari berkas DOCX."

        if nama.endswith(".doc"):
            return "", ("❌ Format .doc (Word lama) belum didukung. Mohon simpan ulang berkas sebagai "
                        ".docx terlebih dahulu (Save As > Word Document .docx) lalu unggah kembali.")

        if nama.endswith(".txt"):
            teks = data.decode("utf-8", errors="ignore").strip()
            if not teks:
                return "", "⚠️ Berkas TXT kosong."
            return teks[:max_chars], f"✅ Berhasil membaca {len(teks)} karakter dari berkas TXT."

        if nama.endswith(".csv") or nama.endswith(".xlsx") or nama.endswith(".xls"):
            import pandas as pd
            if nama.endswith(".csv"):
                df = pd.read_csv(BytesIO(data))
            else:
                df = pd.read_excel(BytesIO(data))
            if df.empty:
                return "", "⚠️ Berkas tabel terbaca tapi tidak ada data di dalamnya."
            teks = df.to_string(index=False)
            return teks[:max_chars], f"✅ Berhasil membaca {len(df)} baris data dari berkas tabel."

        if nama.endswith((".jpg", ".jpeg", ".png")):
            return "", ("⚠️ Berkas berupa gambar. AI teks tidak dapat membaca isi gambar secara otomatis — "
                        "gambar tetap bisa dipakai sebagai ilustrasi (jika fitur ini menyediakan unggah gambar), "
                        "namun isinya tidak dijadikan konteks teks.")
    except Exception as e:
        return "", f"❌ Gagal mengekstrak isi berkas ({uploaded_file.name}): {e}"

    return "", "⚠️ Format berkas tidak didukung."


def uploader_referensi(key: str, label: str = "📎 Unggah Referensi Tambahan (Buku Guru / Materi - PDF, DOCX, TXT, opsional)",
                        tipe: list = None):
    """Widget unggah file referensi yang dipakai untuk memperkaya konteks prompt AI.
    Menampilkan status pembacaan secara jelas + cuplikan isi agar guru bisa memastikan berkas terbaca."""
    tipe = tipe or ["pdf", "docx", "txt"]
    f = st.file_uploader(label, type=tipe, key=key)
    if f is None:
        return ""
    with st.spinner(f"Membaca isi berkas {f.name}..."):
        teks, status = ekstrak_teks_referensi(f)
    if teks:
        st.success(status)
        with st.expander(f"👀 Pratinjau isi yang terbaca dari '{f.name}' (klik untuk lihat)"):
            st.text(teks[:2000] + ("..." if len(teks) > 2000 else ""))
        return teks
    else:
        st.warning(status or f"⚠️ Tidak dapat membaca isi berkas '{f.name}'.")
        return ""


def baca_tabel_upload(uploaded_file):
    """Baca berkas CSV/XLSX menjadi list of dict (baris), dipakai untuk impor massal
    data siswa / jurnal / kaldik. Mengembalikan (list_baris, pesan_status)."""
    if uploaded_file is None:
        return [], ""
    nama = uploaded_file.name.lower()
    try:
        import pandas as pd
        data = uploaded_file.getvalue()
        if nama.endswith(".csv"):
            df = pd.read_csv(BytesIO(data))
        elif nama.endswith(".xlsx") or nama.endswith(".xls"):
            df = pd.read_excel(BytesIO(data))
        elif nama.endswith(".txt"):
            teks = data.decode("utf-8", errors="ignore")
            baris = [{"Nama": b.strip()} for b in teks.split("\n") if b.strip()]
            return baris, f"✅ {len(baris)} baris terbaca dari TXT."
        else:
            return [], "⚠️ Format tidak didukung untuk impor tabel (gunakan CSV/XLSX/TXT)."
        df = df.fillna("")
        baris = df.to_dict(orient="records")
        return baris, f"✅ {len(baris)} baris data berhasil dibaca dari '{uploaded_file.name}'."
    except Exception as e:
        return [], f"❌ Gagal membaca berkas '{uploaded_file.name}': {e}"


# ============================================================
# FUNGSI PEMANGGILAN GROQ
# ============================================================
def bersihkan_html(teks: str) -> str:
    """Jaring pengaman: kalau model AI tetap menyelipkan tag HTML (misal <textarea>, <img>,
    <br>) meski sudah dilarang di prompt, buang tag-tag itu supaya tidak muncul sebagai
    teks aneh/rusak di dokumen PDF/Word."""
    if not teks:
        return teks
    # Buang tag HTML utuh (misal <img src="...">) beserta pasangannya kalau ada
    teks = re.sub(r"<(textarea|div|input|table|tr|td|th|thead|tbody)[^>]*>.*?</\1>",
                  " ______________ ", teks, flags=re.IGNORECASE | re.DOTALL)
    # Ganti <br> jadi baris baru
    teks = re.sub(r"<br\s*/?>", "\n", teks, flags=re.IGNORECASE)
    # Buang sisa tag HTML tunggal apa pun yang masih tersisa (termasuk <img ...>)
    teks = re.sub(r"<[^>]+>", "", teks)
    return teks


def call_groq(api_key: str, prompt: str, max_tokens: int = 6000) -> str:
    """Panggil Groq. Kalau kena limit token-per-menit (umum di akun gratis, sekarang
    dibatasi ketat ~8.000 TPM untuk hampir semua model), otomatis coba lagi dengan
    permintaan token keluaran yang lebih kecil, sampai 3x percobaan."""
    client = Groq(api_key=api_key)
    coba_max = max_tokens
    error_terakhir = None
    for percobaan in range(3):
        try:
            completion = client.chat.completions.create(
                model=MODEL_GROQ,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.6,
                max_completion_tokens=coba_max,
            )
            return bersihkan_html(completion.choices[0].message.content)
        except Exception as e:
            error_terakhir = e
            pesan = str(e)
            if "rate_limit_exceeded" in pesan or "tokens per minute" in pesan or "Request too large" in pesan:
                coba_max = max(1500, coba_max // 2)
                continue
            raise
    raise RuntimeError(
        "Permintaan ke Groq berulang kali ditolak karena melebihi batas token per menit "
        "(akun API Key gratis Groq sekarang dibatasi ketat, sekitar 8.000 token/menit). "
        "Coba: (1) tunggu 1 menit lalu ulangi, (2) persingkat isian/materi yang dimasukkan, "
        "atau (3) upgrade API Key ke Groq Dev Tier di console.groq.com/settings/billing "
        f"(biayanya sangat murah). Detail teknis: {error_terakhir}"
    )


# ============================================================
# PARSER MARKDOWN BERSAMA (dipakai untuk PDF & DOCX)
# ============================================================
def parse_markdown_blocks(md_text: str) -> list:
    """Mengubah teks markdown dari LLM menjadi daftar blok terstruktur:
    ('h1'|'h2'|'h3', text) / ('bullet'|'numbered'|'para', text) / ('table', rows)"""
    blocks = []
    lines = md_text.replace("\r\n", "\n").split("\n")
    i, n = 0, len(lines)

    while i < n:
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith("|"):
            table_rows = []
            while i < n and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if not table_rows and i + 1 < n and re.match(r"^\|?[\s:\-|]+\|?$", lines[i + 1].strip()):
                    i += 2
                    table_rows.append([c.strip() for c in row_line.strip("|").split("|")])
                    continue
                table_rows.append([c.strip() for c in row_line.strip("|").split("|")])
                i += 1
            if table_rows:
                blocks.append(("table", table_rows))
            continue

        if line.startswith("# "):
            blocks.append(("h1", line[2:])); i += 1; continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:])); i += 1; continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:])); i += 1; continue
        if re.match(r"^[-*]\s+", line):
            blocks.append(("bullet", re.sub(r"^[-*]\s+", "", line))); i += 1; continue
        if re.match(r"^\d+[.)]\s+", line):
            blocks.append(("numbered", re.sub(r"^\d+[.)]\s+", "", line))); i += 1; continue

        blocks.append(("para", line)); i += 1

    return blocks


# ---------- Rendering ke PDF (reportlab) ----------
def _inline_format_pdf(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)", r"<i>\1</i>", text)
    return text


def build_pdf_styles():
    base = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle("H1", parent=base["Heading1"], fontSize=15,
                              textColor=COLOR_PRIMARY, spaceAfter=4, alignment=TA_LEFT),
        "h2": ParagraphStyle("H2", parent=base["Heading2"], fontSize=12.5,
                              textColor=COLOR_PRIMARY, spaceAfter=4, alignment=TA_LEFT),
        "h3": ParagraphStyle("H3", parent=base["Heading3"], fontSize=11,
                              textColor=colors.HexColor("#333333"), spaceAfter=3),
        "body": ParagraphStyle("Body", parent=base["Normal"], fontSize=10.3,
                                leading=15, alignment=TA_JUSTIFY, spaceAfter=5),
        "bullet": ParagraphStyle("Bullet", parent=base["Normal"], fontSize=10.3,
                                  leading=14, leftIndent=14, spaceAfter=3, alignment=TA_JUSTIFY),
        "cell": ParagraphStyle("Cell", parent=base["Normal"], fontSize=9.3, leading=12),
        "meta_label": ParagraphStyle("MetaLabel", parent=base["Normal"], fontSize=9.5,
                                      textColor=colors.HexColor("#555555"), alignment=TA_RIGHT),
        "meta_value": ParagraphStyle("MetaValue", parent=base["Normal"], fontSize=10,
                                      fontName="Helvetica-Bold"),
        "title": ParagraphStyle("DocTitle", parent=base["Heading1"], fontSize=18,
                                 alignment=TA_CENTER, textColor=COLOR_PRIMARY, spaceAfter=2),
        "subtitle": ParagraphStyle("DocSubtitle", parent=base["Normal"], fontSize=10.5,
                                    alignment=TA_CENTER, textColor=colors.HexColor("#555555"),
                                    spaceAfter=10),
    }


def blocks_to_pdf_flowables(blocks: list, styles: dict) -> list:
    flowables = []
    for kind, content in blocks:
        if kind == "table":
            para_rows = [[Paragraph(_inline_format_pdf(c), styles["cell"]) for c in row] for row in content]
            tbl = Table(para_rows, repeatRows=1, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), COLOR_PRIMARY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, COLOR_LIGHT]),
                ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            flowables.append(tbl)
            flowables.append(Spacer(1, 10))
        elif kind == "h1":
            flowables.append(Spacer(1, 6))
            flowables.append(Paragraph(_inline_format_pdf(content), styles["h1"]))
            flowables.append(HRFlowable(width="100%", thickness=1.2, color=COLOR_PRIMARY, spaceAfter=8))
        elif kind == "h2":
            flowables.append(Spacer(1, 8))
            flowables.append(Paragraph(_inline_format_pdf(content), styles["h2"]))
        elif kind == "h3":
            flowables.append(Spacer(1, 4))
            flowables.append(Paragraph(_inline_format_pdf(content), styles["h3"]))
        elif kind == "bullet":
            flowables.append(Paragraph("• " + _inline_format_pdf(content), styles["bullet"]))
        elif kind == "numbered":
            flowables.append(Paragraph(_inline_format_pdf(content), styles["bullet"]))
        else:
            flowables.append(Paragraph(_inline_format_pdf(content), styles["body"]))
    return flowables


def _pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(COLOR_ACCENT)
    canvas.setLineWidth(2)
    canvas.line(2 * cm, 1.6 * cm, A4[0] - 2 * cm, 1.6 * cm)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#777777"))
    canvas.drawString(2 * cm, 1.2 * cm,
                       f"Dibuat dengan Generator Perangkat Ajar KKG - {datetime.now().strftime('%d %B %Y')}")
    canvas.drawRightString(A4[0] - 2 * cm, 1.2 * cm, f"Halaman {doc.page}")
    canvas.restoreState()


def _meta_pairs(meta: dict):
    rows, pair = [], []
    for k, v in meta.items():
        if not v:
            continue
        pair.append((k, v))
        if len(pair) == 2:
            rows.append(pair); pair = []
    if pair:
        rows.append(pair)
    return rows


def generate_pdf(doc_title: str, meta: dict, markdown_text: str, gambar_bytes: bytes = None) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=1.8 * cm, bottomMargin=2.1 * cm, leftMargin=2 * cm, rightMargin=2 * cm,
        title=doc_title,
    )
    styles = build_pdf_styles()
    story = [
        Paragraph(doc_title, styles["title"]),
        Paragraph("Kurikulum Merdeka &bull; Pembelajaran Mendalam (Deep Learning)", styles["subtitle"]),
    ]

    rows = _meta_pairs(meta)
    if rows:
        table_data = []
        for pair in rows:
            row = []
            for k, v in pair:
                row.append(Paragraph(f"<b>{k}</b>", styles["meta_label"]))
                row.append(Paragraph(f": {v}", styles["meta_value"]))
            if len(pair) == 1:
                row += ["", ""]
            table_data.append(row)
        # Label rata-kanan pada kolom lebar tetap -> tanda titik dua otomatis lurus,
        # tanpa perlu kolom sempit terpisah (yang rawan error karena lebih kecil dari padding sel).
        # Total HARUS <= lebar halaman - margin kiri/kanan (A4 21cm - 2cm - 2cm = 17cm).
        meta_table = Table(table_data, colWidths=[3.4 * cm, 5.4 * cm, 3.0 * cm, 5.0 * cm])
        meta_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), COLOR_LIGHT),
            ("BOX", (0, 0), (-1, -1), 0.7, COLOR_PRIMARY),
            ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.white),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 12))

    if gambar_bytes:
        try:
            img_buf = BytesIO(gambar_bytes)
            rl_img = RLImage(img_buf, width=8 * cm, height=6 * cm, kind="proportional")
            rl_img.hAlign = "CENTER"
            story.append(rl_img)
            story.append(Paragraph("Gambar Ilustrasi Stimulus Soal", styles["subtitle"]))
            story.append(Spacer(1, 8))
        except Exception:
            pass

    story.extend(blocks_to_pdf_flowables(parse_markdown_blocks(markdown_text), styles))
    doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------- Rendering ke DOCX (python-docx, bisa diedit) ----------
def _add_formatted_runs_docx(paragraph, text: str):
    tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2]); run.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1]); run.italic = True
        else:
            paragraph.add_run(tok)


def generate_docx(doc_title: str, meta: dict, markdown_text: str, gambar_bytes: bytes = None) -> bytes:
    document = Document()

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title)
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = DOCX_PRIMARY_RGB

    p_sub = document.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Kurikulum Merdeka • Pembelajaran Mendalam (Deep Learning)")
    r_sub.italic = True
    r_sub.font.size = Pt(10)
    r_sub.font.color.rgb = DOCX_GREY_RGB

    rows = _meta_pairs(meta)
    if rows:
        table = document.add_table(rows=0, cols=4)
        table.style = "Light Grid Accent 1"
        table.autofit = False
        table.allow_autofit = False
        # Label rata-kanan pada kolom lebar tetap -> titik dua otomatis lurus di setiap baris.
        # Total <= lebar halaman - margin kiri/kanan (21cm - 2cm - 2cm = 17cm)
        lebar_kolom = [Cm(3.4), Cm(5.4), Cm(3.0), Cm(5.0)]
        for pair in rows:
            cells = table.add_row().cells
            for c, w in zip(cells, lebar_kolom):
                c.width = w
            idx = 0
            for k, v in pair:
                cells[idx].text = ""
                run = cells[idx].paragraphs[0].add_run(k)
                run.bold = True
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells[idx + 1].text = f": {v}"
                idx += 2
            if len(pair) == 1:
                for extra_idx in range(2, 4):
                    cells[extra_idx].text = ""
        document.add_paragraph("")

    if gambar_bytes:
        try:
            document.add_picture(BytesIO(gambar_bytes), width=Inches(3.2))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = document.add_paragraph("Gambar Ilustrasi Stimulus Soal")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
        except Exception:
            pass

    for kind, content in parse_markdown_blocks(markdown_text):
        if kind == "h1":
            h = document.add_heading(level=1)
            _add_formatted_runs_docx(h, content)
        elif kind == "h2":
            h = document.add_heading(level=2)
            _add_formatted_runs_docx(h, content)
        elif kind == "h3":
            h = document.add_heading(level=3)
            _add_formatted_runs_docx(h, content)
        elif kind == "bullet":
            p = document.add_paragraph(style="List Bullet")
            _add_formatted_runs_docx(p, content)
        elif kind == "numbered":
            p = document.add_paragraph(style="List Number")
            _add_formatted_runs_docx(p, content)
        elif kind == "table":
            ncols = max(len(r) for r in content)
            tbl = document.add_table(rows=0, cols=ncols)
            tbl.style = "Light Grid Accent 1"
            for ridx, row in enumerate(content):
                cells = tbl.add_row().cells
                for cidx in range(ncols):
                    text = row[cidx] if cidx < len(row) else ""
                    cells[cidx].text = ""
                    _add_formatted_runs_docx(cells[cidx].paragraphs[0], text)
                    if ridx == 0:
                        for run in cells[cidx].paragraphs[0].runs:
                            run.bold = True
            document.add_paragraph("")
        else:
            p = document.add_paragraph()
            _add_formatted_runs_docx(p, content)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def download_row(base_filename: str, pdf_bytes: bytes, docx_bytes: bytes, key_prefix: str):
    c1, c2 = st.columns(2)
    with c1:
        st.download_button("⬇️ Download PDF (A4, cetak)", data=pdf_bytes,
                            file_name=f"{base_filename}.pdf", mime="application/pdf",
                            key=f"{key_prefix}_pdf")
    with c2:
        st.download_button(
            "⬇️ Download DOC (bisa diedit)", data=docx_bytes,
            file_name=f"{base_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"{key_prefix}_docx",
        )


# ============================================================
# SIDEBAR - API KEY & IDENTITAS SEKOLAH (dipakai lintas tab)
# ============================================================
with st.sidebar:
    if AUTH_AKTIF and st.session_state.user:
        st.success(f"👤 Masuk sebagai: **{st.session_state.user['email']}**")
        if st.button("🚪 Keluar", key="btn_keluar", use_container_width=True):
            sb_keluar()
            st.rerun()
        st.divider()
    st.header("⚙️ Pengaturan")
    groq_api_key = st.text_input(
        "🔑 API Key Groq", type="password",
        help="Dapatkan API key gratis di console.groq.com. Key TIDAK disimpan di server, "
             "hanya dipakai selama sesi ini.",
    )
    st.caption(
        "⚠️ Jangan pernah menempelkan API key langsung di dalam kode. Jika key Anda pernah "
        "tertulis di source code / dibagikan publik, segera regenerate di console.groq.com."
    )
    st.divider()
    st.subheader("🏫 Identitas Sekolah")
    st.session_state.sekolah = st.text_input("Nama Sekolah", st.session_state.sekolah)
    st.session_state.tahun_ajaran = st.text_input("Tahun Ajaran", st.session_state.tahun_ajaran)
    st.divider()
    st.subheader("👩‍🏫 Identitas Guru Penyusun")
    st.session_state.penyusun = st.text_input("Nama Guru Penyusun", st.session_state.penyusun)
    st.session_state.guru_nip = st.text_input("NIP Guru Penyusun", st.session_state.guru_nip,
                                               placeholder="Contoh: 199001012019031001")
    st.divider()
    st.subheader("🧑‍💼 Identitas Kepala Sekolah")
    st.session_state.kepsek_nama = st.text_input("Nama Kepala Sekolah", st.session_state.kepsek_nama)
    st.session_state.kepsek_nip = st.text_input("NIP Kepala Sekolah", st.session_state.kepsek_nip,
                                                 placeholder="Contoh: 198005052008011003")
    st.divider()
    st.caption("📘 Seluruh materi dihasilkan berpedoman pada Permendikdasmen No. 13 Tahun 2025 "
               "tentang Pembelajaran Mendalam (Mindful, Meaningful, Joyful).")
    st.divider()
    st.caption("🔗 Ingin aplikasi ini online & bisa dibagikan/dijual ke rekan guru lain? "
               "Lihat panduan **README_DEPLOY.md** yang disertakan.")


st.title("📚 Generator Perangkat Ajar Kurikulum Merdeka")
st.caption(f"Aplikasi Pembantu Guru {st.session_state.sekolah} — Powered by Groq / Llama-3.3 — "
           f"Berpedoman Pembelajaran Mendalam (Deep Learning)")

tab_tp, tab_modul, tab_lkpd, tab_prota, tab_minggu, tab_jurnal, tab_absen = st.tabs(
    ["🎯 TP & ATP", "📖 Modul Ajar", "📝 LKPD", "📅 PROTA & PROMES",
     "🗓️ Minggu Efektif", "📔 Jurnal Mengajar", "🧑‍🤝‍🧑 Absensi Siswa"]
)

# ============================================================
# TAB 1 - TUJUAN PEMBELAJARAN (TP) & ALUR TUJUAN PEMBELAJARAN (ATP)
# ============================================================
with tab_tp:
    st.subheader("Rumuskan TP & Susun ATP")
    c1, c2 = st.columns(2)
    with c1:
        tp_mapel = st.selectbox("Mata Pelajaran", SUBJECT_OPTIONS, key="tp_mapel")
        tp_kelas = st.selectbox("Kelas (1-6)", KELAS_OPTIONS, key="tp_kelas")
        tp_semester = st.selectbox("Semester", ["Ganjil", "Genap"], key="tp_semester")
    with c2:
        tp_cp = st.text_area(
            "Capaian Pembelajaran (CP) / Elemen",
            placeholder="Tempel CP resmi dari Kemendikbudristek untuk elemen ini, "
                        "atau tuliskan ringkasan elemen materi.",
            height=120, key="tp_cp",
        )
        tp_jumlah_tp = st.number_input("Perkiraan Jumlah TP yang diinginkan", 3, 20, 6)

    tp_referensi = uploader_referensi("tp_referensi_upl")

    if st.button("🚀 Buat TP & ATP", key="btn_tp"):
        if not groq_api_key:
            st.error("⚠️ Mohon masukkan API Key Groq di sidebar terlebih dahulu.")
        elif not tp_cp:
            st.error("⚠️ Mohon isi Capaian Pembelajaran (CP) / Elemen.")
        else:
            try:
                with st.spinner("Menyusun Tujuan Pembelajaran dan Alur Tujuan Pembelajaran..."):
                    prompt = f"""
Anda adalah Pakar Pengembang Kurikulum Merdeka dan Pengawas Sekolah Senior di Indonesia.
Susun TUJUAN PEMBELAJARAN (TP) dan ALUR TUJUAN PEMBELAJARAN (ATP) yang siap pakai.

{DEEP_LEARNING_GUIDE}

DATA INPUT:
- Satuan Pendidikan: {st.session_state.sekolah}
- Mata Pelajaran: {tp_mapel}
- Kelas: {label_kelas_fase(tp_kelas)}
- Semester: {tp_semester}
- Tahun Ajaran: {st.session_state.tahun_ajaran}
- Capaian Pembelajaran (CP) / Elemen: {tp_cp}
- Target jumlah TP: sekitar {tp_jumlah_tp} tujuan
{f"- REFERENSI TAMBAHAN dari guru (jadikan acuan konteks/istilah bila relevan): {tp_referensi}" if tp_referensi else ""}

FORMAT OUTPUT (Markdown, WAJIB rapi, gunakan tabel markdown dengan simbol | untuk ATP):

# TUJUAN PEMBELAJARAN & ALUR TUJUAN PEMBELAJARAN

## A. IDENTITAS
Sebutkan satuan pendidikan, mapel, kelas/fase, semester, tahun ajaran.

## B. CAPAIAN PEMBELAJARAN (CP)
Tuliskan ulang CP secara ringkas dan jelas.

## C. TUJUAN PEMBELAJARAN (TP)
Buat daftar TP dengan kode (TP 1, TP 2, dst), masing-masing memuat KOMPETENSI + LINGKUP MATERI,
disusun terukur dan operasional (gunakan kata kerja operasional Taksonomi Bloom revisi).

## D. ALUR TUJUAN PEMBELAJARAN (ATP)
Sajikan dalam bentuk TABEL markdown dengan kolom:
| No | Kode TP | Tujuan Pembelajaran | Alokasi Waktu (JP) | Dimensi Profil Lulusan Terkait | Prinsip Deep Learning (Mindful/Meaningful/Joyful) |
Urutkan dari kompetensi prasyarat termudah menuju paling kompleks (logis & berjenjang).

## E. CATATAN PENGEMBANGAN
Berikan catatan singkat strategi asesmen dan diferensiasi yang disarankan sepanjang alur ini.

Gunakan Bahasa Indonesia resmi, komunikatif, dan terstruktur rapi.
"""
                    hasil = call_groq(groq_api_key, prompt)
                    st.session_state.hasil_tp_atp = hasil
                    st.session_state.meta_tp_atp = {
                        "Satuan Pendidikan": st.session_state.sekolah,
                        "Mata Pelajaran": tp_mapel,
                        "Kelas": label_kelas_fase(tp_kelas),
                        "Semester": tp_semester,
                        "Tahun Ajaran": st.session_state.tahun_ajaran,
                        "Guru Penyusun": st.session_state.penyusun,
                        "NIP Guru": st.session_state.guru_nip,
                        "Kepala Sekolah": st.session_state.kepsek_nama,
                        "NIP Kepala Sekolah": st.session_state.kepsek_nip,
                    }
                st.success("✨ TP & ATP berhasil disusun!")
            except Exception as e:
                st.error(f"❌ Terjadi kesalahan API Groq: {e}")

    if st.session_state.hasil_tp_atp:
        st.markdown(st.session_state.hasil_tp_atp)
        pdf_bytes = generate_pdf("TUJUAN PEMBELAJARAN & ALUR TUJUAN PEMBELAJARAN",
                                  st.session_state.meta_tp_atp, st.session_state.hasil_tp_atp)
        docx_bytes = generate_docx("TUJUAN PEMBELAJARAN & ALUR TUJUAN PEMBELAJARAN",
                                    st.session_state.meta_tp_atp, st.session_state.hasil_tp_atp)
        download_row(f"TP_ATP_{tp_mapel.replace(' ', '_')}", pdf_bytes, docx_bytes, "dl_tp")

# ============================================================
# TAB 2 - MODUL AJAR
# ============================================================
with tab_modul:
    st.subheader("Susun Modul Ajar Lengkap")
    c1, c2 = st.columns(2)
    with c1:
        m_mapel = st.selectbox("Mata Pelajaran", SUBJECT_OPTIONS, key="m_mapel")
        m_kelas = st.selectbox("Kelas (1-6)", KELAS_OPTIONS, key="m_kelas")
        m_alokasi = st.text_input("Alokasi Waktu", "2 x 35 Menit (1 Pertemuan)", key="m_alokasi")
    with c2:
        m_topik = st.text_input("Topik / Materi Utama", key="m_topik")
        m_tp_acuan = st.text_area(
            "Tujuan Pembelajaran acuan (opsional, bisa salin dari tab TP & ATP)",
            height=90, key="m_tp_acuan",
        )

    m_referensi = uploader_referensi("m_referensi_upl")

    if st.button("🚀 Buat Modul Ajar Sekarang", key="btn_modul"):
        if not groq_api_key:
            st.error("⚠️ Mohon masukkan API Key Groq di sidebar terlebih dahulu.")
        elif not m_topik:
            st.error("⚠️ Mohon isi Topik / Materi Utama!")
        else:
            try:
                with st.spinner("Groq AI sedang menyusun Modul Ajar presisi... Mohon tunggu..."):
                    prompt = f"""
Anda adalah Pakar Pengembang Kurikulum Merdeka dan Pengawas Sekolah Senior di Indonesia.
Tugas Anda adalah membuatkan MODUL AJAR LENGKAP & SIAP PAKAI sesuai format resmi Kemendikbudristek.

{DEEP_LEARNING_GUIDE}

DATA INPUT:
- Satuan Pendidikan: {st.session_state.sekolah}
- Guru Penyusun: {st.session_state.penyusun or '(diisi oleh guru)'}
- Mata Pelajaran: {m_mapel}
- Kelas: {label_kelas_fase(m_kelas)}
- Topik Utama: {m_topik}
- Alokasi Waktu: {m_alokasi}
- Tujuan Pembelajaran acuan (jika ada): {m_tp_acuan or 'Rumuskan sendiri berdasarkan topik.'}
{f"- REFERENSI TAMBAHAN dari Buku Guru/materi yang diunggah (jadikan acuan isi & istilah bila relevan): {m_referensi}" if m_referensi else ""}

PETUNJUK FORMAT OUTPUT (Wajib Rapi dengan Markdown):

# MODUL AJAR KURIKULUM MERDEKA

## A. INFORMASI UMUM
1. **Identitas Modul**: Nama Penyusun, Sekolah, Jenjang SD, Mata Pelajaran, Alokasi Waktu.
2. **Kompetensi Awal**: Pengetahuan dasar yang harus dimiliki murid sebelum materi ini.
3. **Profil Lulusan / Profil Pelajar Pancasila**: Sebutkan dimensi yang relevan.
4. **Sarana & Prasarana**: Alat dan bahan yang dibutuhkan.
5. **Target Peserta Didik**: Reguler / Tipikal.
6. **Model Pembelajaran**: Tatap Muka / Problem Based Learning / Diferensiasi (sesuaikan Deep Learning).

## B. KOMPONEN INTI
1. **Tujuan Pembelajaran (TP)**: Jabarkan secara rinci dan terukur.
2. **Pemahaman Bermakna**: Manfaat materi dalam kehidupan sehari-hari murid.
3. **Pertanyaan Pemantik**: 2-3 pertanyaan pemicu diskusi di awal kelas.
4. **Kegiatan Pembelajaran**:
   - **Pendahuluan (10 Menit)**: Orientasi, Apersepsi, Motivasi. Beri label prinsip Deep Learning.
   - **Kegiatan Inti (50 Menit)**: Langkah konkret pembelajaran diferensiasi/aktif, SETIAP langkah
     diberi label (Mindful)/(Meaningful)/(Joyful) sesuai pedoman di atas.
   - **Penutup (10 Menit)**: Evaluasi, Refleksi (Mindful), Rencana Pertemuan Berikutnya.
5. **Asesmen**: Asesmen Formatif lengkap dengan Rubrik Penilaian dalam TABEL markdown (Skor 1-4).

## C. LAMPIRAN
1. **Lembar Kerja Peserta Didik (LKPD) ringkas**: Buatkan soal/tugas interaktif siap cetak.
2. **Bahan Bacaan Guru & Murid**: Ringkasan materi singkat.

Gunakan Bahasa Indonesia resmi, komunikatif, dan terstruktur rapi.
"""
                    hasil = call_groq(groq_api_key, prompt)
                    st.session_state.hasil_modul = hasil
                    st.session_state.meta_modul = {
                        "Satuan Pendidikan": st.session_state.sekolah,
                        "Mata Pelajaran": m_mapel,
                        "Kelas": label_kelas_fase(m_kelas),
                        "Topik": m_topik,
                        "Alokasi Waktu": m_alokasi,
                        "Guru Penyusun": st.session_state.penyusun,
                        "NIP Guru": st.session_state.guru_nip,
                        "Kepala Sekolah": st.session_state.kepsek_nama,
                        "NIP Kepala Sekolah": st.session_state.kepsek_nip,
                    }
                st.success("✨ Modul Ajar berhasil disusun!")
            except Exception as e:
                st.error(f"❌ Terjadi kesalahan API Groq: {e}")

    if st.session_state.hasil_modul:
        st.markdown(st.session_state.hasil_modul)
        pdf_bytes = generate_pdf("MODUL AJAR KURIKULUM MERDEKA",
                                  st.session_state.meta_modul, st.session_state.hasil_modul)
        docx_bytes = generate_docx("MODUL AJAR KURIKULUM MERDEKA",
                                    st.session_state.meta_modul, st.session_state.hasil_modul)
        download_row(f"Modul_Ajar_{m_mapel.replace(' ', '_') if m_mapel else 'Modul'}",
                     pdf_bytes, docx_bytes, "dl_modul")

# ============================================================
# TAB 3 - LEMBAR KERJA PESERTA DIDIK (LKPD)
# ============================================================
with tab_lkpd:
    st.subheader("Susun LKPD Siap Cetak")
    c1, c2 = st.columns(2)
    with c1:
        l_mapel = st.selectbox("Mata Pelajaran", SUBJECT_OPTIONS, key="l_mapel")
        l_kelas = st.selectbox("Kelas (1-6)", KELAS_OPTIONS, key="l_kelas")
        l_jumlah_soal = st.slider("Jumlah Soal", 3, 15, 8, key="l_jumlah_soal")
    with c2:
        l_topik = st.text_input("Topik / Materi", key="l_topik")
        l_tingkat = st.select_slider(
            "Tingkat Kesulitan",
            options=["Mudah (C1-C2)", "Sedang (C3-C4)", "Menantang / HOTS (C5-C6)"],
            value="Sedang (C3-C4)", key="l_tingkat",
        )
        l_jenis_soal = st.multiselect(
            "Jenis Soal", ["Pilihan Ganda", "Isian Singkat", "Uraian / Essay", "Praktik/Unjuk Kerja"],
            default=["Pilihan Ganda", "Uraian / Essay"], key="l_jenis_soal",
        )

    st.markdown("**🎨 Stimulus Soal** — bantu murid memahami konteks soal agar tidak bingung")
    c3, c4 = st.columns(2)
    with c3:
        l_stimulus = st.selectbox(
            "Jenis Stimulus pada beberapa soal",
            ["Tanpa stimulus", "Cerita singkat / narasi kontekstual",
             "Ilustrasi gambar (deskripsi/petunjuk gambar)", "Cerita + Ilustrasi gambar"],
            key="l_stimulus",
        )
    with c4:
        l_gambar_upl = st.file_uploader(
            "🖼️ Unggah Gambar Ilustrasi (opsional, akan disisipkan di LKPD)",
            type=["png", "jpg", "jpeg"], key="l_gambar_upl",
        )
        st.session_state.lkpd_gambar_bytes = l_gambar_upl.getvalue() if l_gambar_upl is not None else None

    l_referensi = uploader_referensi("l_referensi_upl", "📎 Unggah Referensi Materi/Buku Guru (opsional)")

    if st.button("🚀 Buat LKPD Sekarang", key="btn_lkpd"):
        if not groq_api_key:
            st.error("⚠️ Mohon masukkan API Key Groq di sidebar terlebih dahulu.")
        elif not l_topik:
            st.error("⚠️ Mohon isi Topik / Materi.")
        else:
            try:
                with st.spinner("Menyusun LKPD siap cetak..."):
                    prompt = f"""
Anda adalah Pakar Pengembang Kurikulum Merdeka di Indonesia.
Susun LEMBAR KERJA PESERTA DIDIK (LKPD) yang siap dicetak untuk murid Sekolah Dasar.

{DEEP_LEARNING_GUIDE}

DATA INPUT:
- Satuan Pendidikan: {st.session_state.sekolah}
- Mata Pelajaran: {l_mapel}
- Kelas: {label_kelas_fase(l_kelas)}
- Topik / Materi: {l_topik}
- Jumlah Soal: {l_jumlah_soal}
- Tingkat Kesulitan: {l_tingkat}
- Jenis Soal yang diminta: {', '.join(l_jenis_soal) if l_jenis_soal else 'campuran sesuai topik'}
- Kebutuhan Stimulus Soal: {l_stimulus}
{f"- REFERENSI TAMBAHAN dari guru (jadikan acuan isi bila relevan): {l_referensi}" if l_referensi else ""}

INSTRUKSI STIMULUS (WAJIB diikuti sesuai kebutuhan stimulus di atas):
- Jika stimulus berupa "Cerita singkat / narasi kontekstual": sisipkan 1 cerita pendek (3-5 kalimat,
  dekat dengan kehidupan anak SD) sebelum sekelompok soal yang berkaitan, agar anak memahami konteks
  sebelum menjawab.
- Jika stimulus berupa "Ilustrasi gambar (deskripsi/petunjuk gambar)": untuk soal yang relevan, tuliskan
  baris "*(Petunjuk ilustrasi untuk guru: [deskripsi singkat gambar yang perlu ditempel/digambar di sini])*"
  tepat sebelum soal tersebut, agar guru tahu gambar apa yang perlu dilampirkan secara manual.
- Jika "Cerita + Ilustrasi gambar": gabungkan keduanya secara proporsional.
- Jika "Tanpa stimulus": langsung ke soal seperti biasa.
Stimulus TIDAK perlu diberikan di setiap nomor, cukup pada beberapa soal yang paling membutuhkan
konteks agar anak tidak bingung.

FORMAT OUTPUT (Markdown, WAJIB rapi):

# LEMBAR KERJA PESERTA DIDIK (LKPD)

## A. IDENTITAS
Mapel, Kelas/Fase, Topik. Sertakan kolom isian: Nama: __________  Kelas: ______  Tanggal: ______

## B. PETUNJUK PENGERJAAN
Instruksi singkat, jelas, ramah anak, sisipkan kalimat pemantik semangat (Joyful).

## C. TUJUAN
1-2 kalimat tujuan belajar dari LKPD ini (Meaningful - kaitkan dengan kehidupan sehari-hari).

## D. KEGIATAN PEMANASAN (Mindful)
1 kegiatan singkat sebelum mengerjakan soal, misal refleksi/pertanyaan pemantik.

## E. SOAL-SOAL
Nomori setiap soal, sesuai jenis soal yang diminta, tingkat kesulitan {l_tingkat}, sediakan
ruang jawaban (garis titik-titik) untuk isian/uraian. Untuk pilihan ganda beri opsi A-D.

## F. KUNCI JAWABAN & PEDOMAN PENSKORAN
Pisahkan bagian ini dengan jelas (halaman guru), sertakan tabel skor per nomor.

Gunakan Bahasa Indonesia yang komunikatif dan ramah anak sesuai jenjang {label_kelas_fase(l_kelas)}.
"""
                    hasil = call_groq(groq_api_key, prompt)
                    st.session_state.hasil_lkpd = hasil
                    st.session_state.meta_lkpd = {
                        "Satuan Pendidikan": st.session_state.sekolah,
                        "Mata Pelajaran": l_mapel,
                        "Kelas": label_kelas_fase(l_kelas),
                        "Topik": l_topik,
                        "Tingkat": l_tingkat,
                        "Guru Penyusun": st.session_state.penyusun,
                        "NIP Guru": st.session_state.guru_nip,
                    }
                st.success("✨ LKPD berhasil disusun!")
            except Exception as e:
                st.error(f"❌ Terjadi kesalahan API Groq: {e}")

    if st.session_state.hasil_lkpd:
        st.markdown(st.session_state.hasil_lkpd)
        if st.session_state.lkpd_gambar_bytes:
            st.image(st.session_state.lkpd_gambar_bytes, caption="Gambar Ilustrasi Stimulus Soal", width=300)
        pdf_bytes = generate_pdf("LEMBAR KERJA PESERTA DIDIK (LKPD)",
                                  st.session_state.meta_lkpd, st.session_state.hasil_lkpd,
                                  gambar_bytes=st.session_state.lkpd_gambar_bytes)
        docx_bytes = generate_docx("LEMBAR KERJA PESERTA DIDIK (LKPD)",
                                    st.session_state.meta_lkpd, st.session_state.hasil_lkpd,
                                    gambar_bytes=st.session_state.lkpd_gambar_bytes)
        download_row(f"LKPD_{l_mapel.replace(' ', '_') if l_mapel else 'LKPD'}",
                     pdf_bytes, docx_bytes, "dl_lkpd")

# ============================================================
# TAB 4 - PROGRAM TAHUNAN (PROTA) & PROGRAM SEMESTER (PROMES)
# ============================================================
with tab_prota:
    st.subheader("Susun Program Tahunan (PROTA) & Program Semester (PROMES)")
    c1, c2 = st.columns(2)
    with c1:
        pp_mapel = st.selectbox("Mata Pelajaran", SUBJECT_OPTIONS, key="pp_mapel")
        pp_kelas = st.selectbox("Kelas (1-6)", KELAS_OPTIONS, key="pp_kelas")
    with c2:
        pp_minggu_ganjil = st.number_input("Minggu Efektif Semester Ganjil", 8, 24, 18, key="pp_mg")
        pp_minggu_genap = st.number_input("Minggu Efektif Semester Genap", 8, 24, 18, key="pp_mgg")

    pp_materi = st.text_area(
        "Daftar Materi Pokok / TP satu tahun (bisa salin dari tab TP & ATP)",
        height=150, key="pp_materi",
        placeholder="Contoh:\nSemester Ganjil: Tata Cara Wudhu, Sholat Fardhu Berjamaah, Adab kepada Orang Tua\n"
                    "Semester Genap: Puasa Ramadhan, Zakat Fitrah, Kisah Nabi",
    )
    pp_referensi = uploader_referensi("pp_referensi_upl",
                                       "📎 Unggah Referensi Tambahan (Buku Guru/Materi/ATP, opsional)")

    colA, colB = st.columns(2)
    with colA:
        gen_prota = st.button("🚀 Buat PROTA", key="btn_prota", use_container_width=True)
    with colB:
        pp_semester_promes = st.selectbox("Semester untuk PROMES", ["Ganjil", "Genap"], key="pp_semester_promes")
        gen_promes = st.button("🚀 Buat PROMES", key="btn_promes", use_container_width=True)

    if gen_prota or gen_promes:
        if not groq_api_key:
            st.error("⚠️ Mohon masukkan API Key Groq di sidebar terlebih dahulu.")
        elif not pp_materi:
            st.error("⚠️ Mohon isi Daftar Materi Pokok / TP satu tahun.")
        else:
            try:
                if gen_prota:
                    judul = "PROGRAM TAHUNAN (PROTA)"
                    with st.spinner("Menyusun Program Tahunan (PROTA)..."):
                        prompt = f"""
Anda adalah Pakar Pengembang Kurikulum Merdeka di Indonesia.
Susun PROGRAM TAHUNAN (PROTA) yang siap pakai untuk satu tahun ajaran penuh (2 semester).

DATA INPUT:
- Satuan Pendidikan: {st.session_state.sekolah}
- Mata Pelajaran: {pp_mapel}
- Kelas: {label_kelas_fase(pp_kelas)}
- Tahun Ajaran: {st.session_state.tahun_ajaran}
- Minggu Efektif Semester Ganjil: {pp_minggu_ganjil} minggu
- Minggu Efektif Semester Genap: {pp_minggu_genap} minggu
- Daftar Materi Pokok / TP setahun: {pp_materi}
{f"- REFERENSI TAMBAHAN dari guru (jadikan acuan bila relevan): {pp_referensi}" if pp_referensi else ""}

FORMAT OUTPUT (Markdown, WAJIB rapi, gunakan tabel markdown dengan simbol |):

# PROGRAM TAHUNAN (PROTA)

## A. IDENTITAS
Satuan pendidikan, mapel, kelas, tahun ajaran.

## B. TABEL PROGRAM TAHUNAN
Sajikan dalam TABEL markdown dengan kolom:
| No | Semester | Materi Pokok / Tujuan Pembelajaran | Alokasi Waktu (JP) | Keterangan |
Kelompokkan baris per semester (Ganjil dahulu, lalu Genap), alokasi waktu total harus
proporsional dengan minggu efektif yang diberikan (asumsikan sekitar 2-4 JP per minggu per materi,
sesuaikan agar realistis).

## C. REKAPITULASI
Tabel ringkas: total JP semester ganjil, total JP semester genap, total JP satu tahun.

Gunakan Bahasa Indonesia resmi dan terstruktur rapi.
"""
                else:
                    judul = f"PROGRAM SEMESTER (PROMES) - SEMESTER {pp_semester_promes.upper()}"
                    minggu_pakai = pp_minggu_ganjil if pp_semester_promes == "Ganjil" else pp_minggu_genap
                    bulan_range = "Juli - Desember" if pp_semester_promes == "Ganjil" else "Januari - Juni"
                    with st.spinner(f"Menyusun Program Semester ({pp_semester_promes})..."):
                        prompt = f"""
Anda adalah Pakar Pengembang Kurikulum Merdeka di Indonesia.
Susun PROGRAM SEMESTER (PROMES) yang siap pakai untuk semester {pp_semester_promes}.

DATA INPUT:
- Satuan Pendidikan: {st.session_state.sekolah}
- Mata Pelajaran: {pp_mapel}
- Kelas: {label_kelas_fase(pp_kelas)}
- Tahun Ajaran: {st.session_state.tahun_ajaran}
- Semester: {pp_semester_promes} (rentang bulan {bulan_range})
- Minggu Efektif: {minggu_pakai} minggu
- Daftar Materi Pokok / TP (ambil bagian yang relevan untuk semester {pp_semester_promes}): {pp_materi}
{f"- REFERENSI TAMBAHAN dari guru (jadikan acuan bila relevan): {pp_referensi}" if pp_referensi else ""}

FORMAT OUTPUT (Markdown, WAJIB rapi, gunakan tabel markdown dengan simbol |):

# PROGRAM SEMESTER (PROMES) - SEMESTER {pp_semester_promes.upper()}

## A. IDENTITAS
Satuan pendidikan, mapel, kelas, tahun ajaran, semester, rentang bulan {bulan_range}.

## B. TABEL PROGRAM SEMESTER
Sajikan dalam TABEL markdown dengan kolom:
| No | Materi Pokok / TP | Alokasi Waktu (JP) | Bulan & Minggu Pelaksanaan | Keterangan |
Sebar materi secara realistis ke {minggu_pakai} minggu efektif dalam rentang {bulan_range},
sisakan alokasi untuk Penilaian Tengah Semester (PTS) dan Penilaian Akhir Semester (PAS)
sebagai baris tersendiri di tabel.

## C. CATATAN
Catatan singkat tentang hari libur/jeda semester yang perlu diperhitungkan guru.

Gunakan Bahasa Indonesia resmi dan terstruktur rapi.
"""
                hasil = call_groq(groq_api_key, prompt)
                st.session_state.hasil_prota_promes = hasil
                st.session_state.judul_prota_promes = judul
                st.session_state.meta_prota_promes = {
                    "Satuan Pendidikan": st.session_state.sekolah,
                    "Mata Pelajaran": pp_mapel,
                    "Kelas": label_kelas_fase(pp_kelas),
                    "Tahun Ajaran": st.session_state.tahun_ajaran,
                    "Guru Penyusun": st.session_state.penyusun,
                    "NIP Guru": st.session_state.guru_nip,
                    "Kepala Sekolah": st.session_state.kepsek_nama,
                    "NIP Kepala Sekolah": st.session_state.kepsek_nip,
                }
                st.success(f"✨ {judul} berhasil disusun!")
            except Exception as e:
                st.error(f"❌ Terjadi kesalahan API Groq: {e}")

    if st.session_state.hasil_prota_promes:
        st.markdown(st.session_state.hasil_prota_promes)
        pdf_bytes = generate_pdf(st.session_state.judul_prota_promes,
                                  st.session_state.meta_prota_promes, st.session_state.hasil_prota_promes)
        docx_bytes = generate_docx(st.session_state.judul_prota_promes,
                                    st.session_state.meta_prota_promes, st.session_state.hasil_prota_promes)
        nama_file = st.session_state.judul_prota_promes.split(" - ")[0].replace(" ", "_")
        download_row(f"{nama_file}_{pp_mapel.replace(' ', '_') if pp_mapel else 'Program'}",
                     pdf_bytes, docx_bytes, "dl_prota")

# ============================================================
# TAB 5 - ANALISIS MINGGU EFEKTIF (berbasis Kalender Pendidikan/KALDIK)
# ============================================================
with tab_minggu:
    st.subheader("Analisis Minggu Efektif Berdasarkan Kalender Pendidikan (KALDIK)")
    st.caption("Hitung otomatis jumlah minggu efektif dalam satu semester berdasarkan rentang semester "
               "dan daftar hari/minggu tidak efektif (libur, PTS/PAS, jeda semester, dll) dari KALDIK sekolah.")

    c1, c2 = st.columns(2)
    with c1:
        me_semester = st.selectbox("Semester", ["Ganjil", "Genap"], key="me_semester")
        me_tanggal_mulai = st.date_input("Tanggal Mulai Semester", value=date(2025, 7, 14), key="me_mulai")
    with c2:
        me_mapel = st.selectbox("Mata Pelajaran (opsional, untuk judul dokumen)", SUBJECT_OPTIONS, key="me_mapel")
        me_tanggal_selesai = st.date_input("Tanggal Selesai Semester", value=date(2025, 12, 20), key="me_selesai")

    st.markdown("**📋 Daftar Minggu/Periode Tidak Efektif (dari KALDIK)** — satu baris per periode, format:\n"
                "`tanggal_awal - tanggal_akhir | keterangan` (format tanggal: YYYY-MM-DD)")

    me_file_kaldik = st.file_uploader(
        "📎 Atau unggah berkas Kaldik (CSV/XLSX dengan kolom tanggal_awal, tanggal_akhir, keterangan — "
        "atau PDF/DOCX/TXT berisi teks kaldik untuk dibaca sebagai referensi)",
        type=["csv", "xlsx", "xls", "pdf", "docx", "txt"], key="me_file_kaldik",
    )
    if me_file_kaldik is not None:
        nama_kaldik = me_file_kaldik.name.lower()
        if nama_kaldik.endswith((".csv", ".xlsx", ".xls")):
            baris_kaldik, status_kaldik = baca_tabel_upload(me_file_kaldik)
            if baris_kaldik:
                st.success(status_kaldik)
                with st.expander("👀 Pratinjau data kaldik yang terbaca"):
                    st.dataframe(baris_kaldik, use_container_width=True, hide_index=True)
                if st.button("➕ Tambahkan Baris di Atas ke Kolom Kaldik", key="btn_tambah_kaldik_upload"):
                    baris_baru = []
                    for r in baris_kaldik:
                        keys_lower = {str(k).strip().lower(): k for k in r.keys()}
                        k_awal = next((keys_lower[k] for k in keys_lower if "awal" in k or "mulai" in k), None)
                        k_akhir = next((keys_lower[k] for k in keys_lower if "akhir" in k or "selesai" in k), None)
                        k_ket = next((keys_lower[k] for k in keys_lower if "ket" in k), None)
                        if k_awal and k_akhir:
                            try:
                                ta = pd.to_datetime(r[k_awal]).strftime("%Y-%m-%d")
                                tb = pd.to_datetime(r[k_akhir]).strftime("%Y-%m-%d")
                                ket = str(r.get(k_ket, "Tidak Efektif")) if k_ket else "Tidak Efektif"
                                baris_baru.append(f"{ta} - {tb} | {ket}")
                            except Exception:
                                continue
                    if baris_baru:
                        gabungan = (st.session_state.me_kaldik_text + "\n" if st.session_state.me_kaldik_text else "")
                        st.session_state.me_kaldik_text = gabungan + "\n".join(baris_baru)
                        st.success(f"✅ {len(baris_baru)} periode ditambahkan ke kolom kaldik di bawah.")
                        st.rerun()
                    else:
                        st.warning("⚠️ Tidak menemukan kolom tanggal_awal/tanggal_akhir yang bisa dikenali di berkas ini.")
            else:
                st.warning(status_kaldik)
        else:
            teks_kaldik, status_kaldik = ekstrak_teks_referensi(me_file_kaldik)
            if teks_kaldik:
                st.success(status_kaldik)
                with st.expander(f"👀 Pratinjau isi '{me_file_kaldik.name}' — salin manual periode liburnya ke kolom di bawah"):
                    st.text(teks_kaldik[:3000])
            else:
                st.warning(status_kaldik)

    me_kaldik_text = st.text_area(
        "Input Kaldik (libur, PTS, PAS, jeda semester, dll)",
        height=150, key="me_kaldik_text",
        placeholder="2025-07-14 - 2025-07-19 | Libur Akhir Semester Genap Sebelumnya\n"
                    "2025-08-17 - 2025-08-17 | Libur HUT RI\n"
                    "2025-10-06 - 2025-10-11 | Penilaian Tengah Semester (PTS)\n"
                    "2025-12-08 - 2025-12-13 | Penilaian Akhir Semester (PAS)\n"
                    "2025-12-15 - 2025-12-20 | Libur Akhir Semester",
    )

    hitung_me = st.button("🧮 Hitung Minggu Efektif", key="btn_hitung_me", use_container_width=True)

    if hitung_me:
        if me_tanggal_selesai <= me_tanggal_mulai:
            st.error("⚠️ Tanggal selesai harus setelah tanggal mulai semester.")
        else:
            # Parsing baris kaldik
            periode_libur = []
            for baris in me_kaldik_text.strip().split("\n"):
                baris = baris.strip()
                if not baris:
                    continue
                try:
                    bagian_tanggal, keterangan = baris.split("|", 1)
                except ValueError:
                    bagian_tanggal, keterangan = baris, "Tidak Efektif"
                m = re.match(r"\s*(\d{4}-\d{2}-\d{2})\s*-\s*(\d{4}-\d{2}-\d{2})\s*", bagian_tanggal)
                if not m:
                    continue
                try:
                    tgl_awal = datetime.strptime(m.group(1), "%Y-%m-%d").date()
                    tgl_akhir = datetime.strptime(m.group(2), "%Y-%m-%d").date()
                    periode_libur.append((tgl_awal, tgl_akhir, keterangan.strip()))
                except Exception:
                    continue

            # Bangun daftar minggu (Senin-Minggu) yang beririsan rentang semester
            hasil_minggu = []
            senin = me_tanggal_mulai - timedelta(days=me_tanggal_mulai.weekday())
            no_minggu = 1
            while senin <= me_tanggal_selesai:
                minggu_akhir = senin + timedelta(days=6)
                keterangan_tidak_efektif = []
                for (la, lb, ket) in periode_libur:
                    if la <= minggu_akhir and lb >= senin:
                        keterangan_tidak_efektif.append(ket)
                status = "Tidak Efektif" if keterangan_tidak_efektif else "Efektif"
                hasil_minggu.append({
                    "No": no_minggu,
                    "Minggu Ke": f"{senin.strftime('%d %b')} - {minggu_akhir.strftime('%d %b %Y')}",
                    "Status": status,
                    "Keterangan": "; ".join(sorted(set(keterangan_tidak_efektif))) if keterangan_tidak_efektif else "-",
                })
                no_minggu += 1
                senin += timedelta(days=7)

            total_minggu = len(hasil_minggu)
            minggu_efektif = sum(1 for m_ in hasil_minggu if m_["Status"] == "Efektif")
            minggu_tidak_efektif = total_minggu - minggu_efektif

            st.session_state.tabel_minggu_efektif = hasil_minggu
            st.session_state.meta_minggu_efektif = {
                "Satuan Pendidikan": st.session_state.sekolah,
                "Semester": me_semester,
                "Tahun Ajaran": st.session_state.tahun_ajaran,
                "Rentang Semester": f"{me_tanggal_mulai.strftime('%d %b %Y')} - {me_tanggal_selesai.strftime('%d %b %Y')}",
                "Guru Penyusun": st.session_state.penyusun,
                "NIP Guru": st.session_state.guru_nip,
            }

            md = f"# ANALISIS MINGGU EFEKTIF SEMESTER {me_semester.upper()}\n\n"
            md += "## A. REKAPITULASI\n\n"
            md += "| Keterangan | Jumlah |\n|---|---|\n"
            md += f"| Total Minggu dalam Semester | {total_minggu} minggu |\n"
            md += f"| Minggu Efektif | {minggu_efektif} minggu |\n"
            md += f"| Minggu Tidak Efektif | {minggu_tidak_efektif} minggu |\n\n"
            md += "## B. RINCIAN PER MINGGU\n\n"
            md += "| No | Minggu Ke | Status | Keterangan |\n|---|---|---|---|\n"
            for m_ in hasil_minggu:
                md += f"| {m_['No']} | {m_['Minggu Ke']} | {m_['Status']} | {m_['Keterangan']} |\n"
            st.session_state.hasil_minggu_efektif = md
            st.success(f"✨ Minggu efektif berhasil dihitung: **{minggu_efektif} dari {total_minggu} minggu** efektif.")

    if st.session_state.hasil_minggu_efektif:
        colm1, colm2, colm3 = st.columns(3)
        tabel = st.session_state.tabel_minggu_efektif or []
        total_ = len(tabel)
        efektif_ = sum(1 for m_ in tabel if m_["Status"] == "Efektif")
        colm1.metric("Total Minggu", total_)
        colm2.metric("Minggu Efektif", efektif_)
        colm3.metric("Minggu Tidak Efektif", total_ - efektif_)
        st.markdown(st.session_state.hasil_minggu_efektif)
        pdf_bytes = generate_pdf(f"ANALISIS MINGGU EFEKTIF - SEMESTER {me_semester.upper()}",
                                  st.session_state.meta_minggu_efektif, st.session_state.hasil_minggu_efektif)
        docx_bytes = generate_docx(f"ANALISIS MINGGU EFEKTIF - SEMESTER {me_semester.upper()}",
                                    st.session_state.meta_minggu_efektif, st.session_state.hasil_minggu_efektif)
        download_row("Analisis_Minggu_Efektif", pdf_bytes, docx_bytes, "dl_me")

# ============================================================
# TAB 6 - JURNAL MENGAJAR HARIAN
# ============================================================
with tab_jurnal:
    st.subheader("Jurnal Mengajar Harian")
    st.caption("Catat kegiatan mengajar setiap hari, lalu unduh rekap jurnal dalam PDF/DOCX kapan saja.")

    with st.form("form_jurnal", clear_on_submit=True):
        jc1, jc2, jc3 = st.columns(3)
        with jc1:
            j_tanggal = st.date_input("Tanggal", value=date.today(), key="j_tanggal")
            j_kelas = st.selectbox("Kelas", KELAS_OPTIONS, key="j_kelas")
        with jc2:
            j_mapel = st.selectbox("Mata Pelajaran", SUBJECT_OPTIONS, key="j_mapel")
            j_jampel = st.text_input("Jam Pelajaran (JP)", "1-2", key="j_jampel")
        with jc3:
            j_hadir = st.text_input("Jumlah Siswa Hadir", "", key="j_hadir")
        j_materi = st.text_input("Materi / Kompetensi yang Diajarkan", key="j_materi")
        j_kegiatan = st.text_area("Ringkasan Kegiatan Pembelajaran", height=80, key="j_kegiatan")
        j_catatan = st.text_area("Catatan / Kendala / Tindak Lanjut (opsional)", height=60, key="j_catatan")
        submit_jurnal = st.form_submit_button("➕ Tambahkan ke Jurnal")

    if submit_jurnal:
        if not j_materi:
            st.error("⚠️ Mohon isi Materi / Kompetensi yang diajarkan.")
        else:
            tambah_baris_jurnal({
                "Tanggal": j_tanggal.strftime("%d-%m-%Y"),
                "Kelas": j_kelas,
                "Mapel": j_mapel,
                "JP": j_jampel,
                "Materi": j_materi,
                "Kegiatan": j_kegiatan,
                "Hadir": j_hadir,
                "Catatan": j_catatan,
            })
            st.success("✅ Catatan jurnal ditambahkan.")

    with st.expander("📎 Atau impor banyak catatan jurnal sekaligus dari berkas CSV/XLSX"):
        st.caption("Kolom yang dikenali: Tanggal, Kelas, Mapel, JP, Materi, Kegiatan, Hadir, Catatan "
                   "(kolom yang tidak ada akan dikosongkan).")
        j_file_upl = st.file_uploader("Unggah berkas jurnal", type=["csv", "xlsx", "xls"], key="j_file_upl")
        if j_file_upl is not None:
            baris_jurnal, status_jurnal = baca_tabel_upload(j_file_upl)
            if baris_jurnal:
                st.success(status_jurnal)
                st.dataframe(baris_jurnal, use_container_width=True, hide_index=True)
                if st.button("➕ Tambahkan Semua ke Jurnal", key="btn_impor_jurnal"):
                    for r in baris_jurnal:
                        keys_lower = {str(k).strip().lower(): k for k in r.keys()}
                        def ambil(*alias):
                            for a in alias:
                                if a in keys_lower:
                                    return str(r[keys_lower[a]])
                            return ""
                        tambah_baris_jurnal({
                            "Tanggal": ambil("tanggal") or date.today().strftime("%d-%m-%Y"),
                            "Kelas": ambil("kelas"),
                            "Mapel": ambil("mapel", "mata pelajaran"),
                            "JP": ambil("jp", "jam pelajaran"),
                            "Materi": ambil("materi"),
                            "Kegiatan": ambil("kegiatan"),
                            "Hadir": ambil("hadir"),
                            "Catatan": ambil("catatan"),
                        })
                    st.success(f"✅ {len(baris_jurnal)} catatan jurnal berhasil diimpor.")
                    st.rerun()
            else:
                st.warning(status_jurnal)

    if st.session_state.jurnal_rows:
        st.markdown(f"**Total catatan jurnal: {len(st.session_state.jurnal_rows)}**")
        st.dataframe([{k: v for k, v in r.items() if k != "id"} for r in st.session_state.jurnal_rows],
                     use_container_width=True, hide_index=True)

        cja, cjb = st.columns(2)
        with cja:
            if st.button("🗑️ Hapus Catatan Terakhir", key="btn_hapus_jurnal"):
                hapus_jurnal_terakhir()
                st.rerun()
        with cjb:
            if st.button("🧹 Kosongkan Semua Jurnal", key="btn_kosong_jurnal"):
                kosongkan_jurnal()
                st.rerun()

        md_jurnal = "# JURNAL MENGAJAR HARIAN\n\n"
        md_jurnal += "| No | Tanggal | Kelas | Mapel | JP | Materi | Hadir | Catatan |\n"
        md_jurnal += "|---|---|---|---|---|---|---|---|\n"
        for i, r in enumerate(st.session_state.jurnal_rows, 1):
            md_jurnal += (f"| {i} | {r['Tanggal']} | {r['Kelas']} | {r['Mapel']} | {r['JP']} | "
                          f"{r['Materi']} | {r['Hadir']} | {r['Catatan'] or '-'} |\n")

        meta_jurnal = {
            "Satuan Pendidikan": st.session_state.sekolah,
            "Tahun Ajaran": st.session_state.tahun_ajaran,
            "Guru Penyusun": st.session_state.penyusun,
            "NIP Guru": st.session_state.guru_nip,
        }
        pdf_bytes = generate_pdf("JURNAL MENGAJAR HARIAN", meta_jurnal, md_jurnal)
        docx_bytes = generate_docx("JURNAL MENGAJAR HARIAN", meta_jurnal, md_jurnal)
        download_row("Jurnal_Mengajar", pdf_bytes, docx_bytes, "dl_jurnal")
    else:
        st.info("Belum ada catatan jurnal. Isi formulir di atas untuk menambahkan.")

# ============================================================
# TAB 7 - DATA SISWA & LEMBAR ABSENSI
# ============================================================
with tab_absen:
    sub_data, sub_kartu, sub_scan, sub_rekap = st.tabs(
        ["📋 Data Siswa & Lembar Absen", "🖨️ Cetak Kartu QR", "📷 Scan Absen", "📊 Rekap Bulanan"]
    )

with sub_data:
    st.subheader("Data Peserta Didik & Lembar Absensi")
    st.caption("Input data siswa satu kali, lalu buat lembar absen siap cetak kapan saja.")

    ac1, ac2 = st.columns(2)
    with ac1:
        a_kelas = st.selectbox("Kelas", KELAS_OPTIONS, key="a_kelas")
    with ac2:
        a_jumlah_pertemuan = st.number_input("Jumlah Kolom Pertemuan/Tanggal pada Lembar Absen", 1, 31, 10, key="a_jp")

    st.markdown("**➕ Tambah Data Siswa**")
    with st.form("form_siswa", clear_on_submit=True):
        sc1, sc2, sc3 = st.columns([3, 1, 1])
        with sc1:
            s_nama = st.text_input("Nama Siswa", key="s_nama")
        with sc2:
            s_jk = st.selectbox("L/P", ["L", "P"], key="s_jk")
        with sc3:
            s_nisn = st.text_input("NISN (opsional)", key="s_nisn")
        submit_siswa = st.form_submit_button("➕ Tambahkan Siswa")

    if submit_siswa:
        if not s_nama:
            st.error("⚠️ Mohon isi nama siswa.")
        else:
            tambah_baris_siswa({"Nama": s_nama, "L/P": s_jk, "NISN": s_nisn})
            st.success(f"✅ {s_nama} ditambahkan ke data siswa {a_kelas}.")

    st.markdown("**📋 Atau tempel banyak nama sekaligus (satu nama per baris)**")
    tempel_siswa = st.text_area("Tempel Daftar Nama", height=100, key="tempel_siswa")
    if st.button("📥 Impor dari Daftar Tempel", key="btn_impor_siswa"):
        nama_list = [n.strip() for n in tempel_siswa.split("\n") if n.strip()]
        for n in nama_list:
            tambah_baris_siswa({"Nama": n, "L/P": "-", "NISN": ""})
        st.success(f"✅ {len(nama_list)} siswa diimpor.")
        st.rerun()

    with st.expander("📎 Atau impor data siswa dari berkas CSV/XLSX/TXT (data dari Dapodik, dll)"):
        st.caption("Kolom yang dikenali: Nama, L/P (atau Jenis Kelamin), NISN. Untuk TXT, satu nama per baris.")
        a_file_upl = st.file_uploader("Unggah berkas data siswa", type=["csv", "xlsx", "xls", "txt"], key="a_file_upl")
        if a_file_upl is not None:
            baris_siswa, status_siswa = baca_tabel_upload(a_file_upl)
            if baris_siswa:
                st.success(status_siswa)
                st.dataframe(baris_siswa, use_container_width=True, hide_index=True)
                if st.button("➕ Tambahkan Semua ke Data Siswa", key="btn_impor_siswa_file"):
                    for r in baris_siswa:
                        keys_lower = {str(k).strip().lower(): k for k in r.keys()}
                        def ambil_s(*alias):
                            for a in alias:
                                if a in keys_lower:
                                    return str(r[keys_lower[a]])
                            return ""
                        nama_ = ambil_s("nama", "nama siswa")
                        if not nama_:
                            continue
                        tambah_baris_siswa({
                            "Nama": nama_,
                            "L/P": ambil_s("l/p", "jk", "jenis kelamin") or "-",
                            "NISN": ambil_s("nisn"),
                        })
                    st.success(f"✅ {len(baris_siswa)} siswa berhasil diimpor.")
                    st.rerun()
            else:
                st.warning(status_siswa)

    if st.session_state.siswa_rows:
        st.markdown(f"**Total siswa terdaftar: {len(st.session_state.siswa_rows)}**")
        st.dataframe([{k: v for k, v in r.items() if k != "id"} for r in st.session_state.siswa_rows],
                     use_container_width=True, hide_index=True)

        sca, scb = st.columns(2)
        with sca:
            if st.button("🗑️ Hapus Siswa Terakhir", key="btn_hapus_siswa"):
                hapus_siswa_terakhir()
                st.rerun()
        with scb:
            if st.button("🧹 Kosongkan Semua Data Siswa", key="btn_kosong_siswa"):
                kosongkan_siswa()
                st.rerun()

        if st.button("🖨️ Buat Lembar Absensi", key="btn_buat_absen", use_container_width=True):
            kolom_tanggal = [f"P{i}" for i in range(1, a_jumlah_pertemuan + 1)]
            header = ["No", "Nama Siswa", "L/P"] + kolom_tanggal
            rows_tabel = [header]
            for i, s in enumerate(st.session_state.siswa_rows, 1):
                rows_tabel.append([str(i), s["Nama"], s["L/P"]] + ["" for _ in kolom_tanggal])

            meta_absen = {
                "Satuan Pendidikan": st.session_state.sekolah,
                "Kelas": a_kelas,
                "Tahun Ajaran": st.session_state.tahun_ajaran,
                "Guru Penyusun": st.session_state.penyusun,
                "NIP Guru": st.session_state.guru_nip,
            }

            # PDF lembar absen (orientasi landscape agar kolom pertemuan muat)
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=landscape(A4),
                                     topMargin=1.5 * cm, bottomMargin=1.8 * cm,
                                     leftMargin=1.5 * cm, rightMargin=1.5 * cm,
                                     title="LEMBAR ABSENSI SISWA")
            styles = build_pdf_styles()
            story = [
                Paragraph("LEMBAR ABSENSI SISWA", styles["title"]),
                Paragraph(f"Kelas {a_kelas} &bull; {st.session_state.sekolah}", styles["subtitle"]),
            ]
            para_rows = [[Paragraph(f"<b>{c}</b>" if r == 0 else c, styles["cell"]) for c in row]
                         for r, row in enumerate(rows_tabel)]
            # Lebar kolom pertemuan menyesuaikan otomatis agar tabel selalu pas di halaman
            # landscape A4 (29.7cm - margin kiri/kanan 1.5cm masing-masing = 26.7cm tersedia),
            # sehingga tidak overflow walau jumlah kolom pertemuan banyak.
            lebar_tersedia = landscape(A4)[0] / cm - 1.5 - 1.5  # cm
            lebar_no, lebar_nama, lebar_lp = 1.0, 5.0, 1.0
            sisa = max(lebar_tersedia - (lebar_no + lebar_nama + lebar_lp), 4.0)
            lebar_per_pertemuan = min(1.0, sisa / max(len(kolom_tanggal), 1))
            col_widths = ([lebar_no * cm, lebar_nama * cm, lebar_lp * cm]
                          + [lebar_per_pertemuan * cm] * len(kolom_tanggal))
            tbl = Table(para_rows, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), COLOR_PRIMARY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, COLOR_LIGHT]),
            ]))
            story.append(tbl)
            doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
            buffer.seek(0)
            pdf_absen = buffer.getvalue()

            # DOCX lembar absen
            document = Document()
            section = document.sections[0]
            section.orientation = 1  # landscape
            section.page_width, section.page_height = Cm(29.7), Cm(21.0)
            section.top_margin = section.bottom_margin = Cm(1.5)
            section.left_margin = section.right_margin = Cm(1.5)
            p_title = document.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("LEMBAR ABSENSI SISWA")
            r_title.bold = True; r_title.font.size = Pt(15); r_title.font.color.rgb = DOCX_PRIMARY_RGB
            p_sub = document.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.add_run(f"Kelas {a_kelas} • {st.session_state.sekolah}").italic = True

            tbl_docx = document.add_table(rows=0, cols=len(header))
            tbl_docx.style = "Light Grid Accent 1"
            tbl_docx.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = tbl_docx.add_row().cells
            for idx, h in enumerate(header):
                hdr_cells[idx].text = ""
                run = hdr_cells[idx].paragraphs[0].add_run(h)
                run.bold = True
            for row in rows_tabel[1:]:
                cells = tbl_docx.add_row().cells
                for idx, v in enumerate(row):
                    cells[idx].text = str(v)

            buffer2 = BytesIO()
            document.save(buffer2)
            buffer2.seek(0)
            docx_absen = buffer2.getvalue()

            st.success("✨ Lembar absensi berhasil dibuat!")
            download_row(f"Absensi_{a_kelas.replace(' ', '_')}", pdf_absen, docx_absen, "dl_absen")
    else:
        st.info("Belum ada data siswa. Tambahkan lewat formulir atau impor daftar tempel di atas.")

with sub_kartu:
    st.subheader("🖨️ Cetak Kartu QR Absen")
    st.caption("Kartu berisi Nama, Kelas, NISN, barcode, dan kotak kosong untuk foto yang "
               "ditempel manual — seperti kartu pelajar biasa.")

    siswa_ada_nisn = [s for s in st.session_state.siswa_rows if str(s.get("NISN", "")).strip()]
    siswa_tanpa_nisn = [s for s in st.session_state.siswa_rows if not str(s.get("NISN", "")).strip()]

    if siswa_tanpa_nisn:
        st.warning(f"⚠️ {len(siswa_tanpa_nisn)} siswa belum punya NISN dan tidak akan tercetak "
                   f"kartunya (QR wajib pakai NISN). Lengkapi dulu di tab 'Data Siswa & Lembar Absen'.")

    if not siswa_ada_nisn:
        st.info("Belum ada siswa dengan NISN terisi.")
    else:
        pilih_semua = st.checkbox("Pilih semua siswa (yang punya NISN)", value=True, key="kartu_pilih_semua")
        nama_opsi = [f"{s['Nama']} — {s.get('NISN', '')}" for s in siswa_ada_nisn]
        default_pilih = nama_opsi if pilih_semua else []
        dipilih = st.multiselect("Pilih siswa yang kartunya mau dicetak", nama_opsi,
                                  default=default_pilih, key="kartu_dipilih")
        siswa_terpilih = [s for s, label in zip(siswa_ada_nisn, nama_opsi) if label in dipilih]

        if st.button("🖨️ Buat PDF Kartu QR", key="btn_buat_kartu", use_container_width=True,
                      disabled=not siswa_terpilih):
            from reportlab.pdfgen import canvas as pdfcanvas
            from reportlab.lib.utils import ImageReader

            LEBAR_KARTU, TINGGI_KARTU = 9.0 * cm, 5.6 * cm
            KOLOM, BARIS = 2, 4
            MARGIN_X, MARGIN_Y = 1.35 * cm, 1.0 * cm
            GAP = 0.3 * cm

            buf = BytesIO()
            c = pdfcanvas.Canvas(buf, pagesize=A4)
            lebar_hal, tinggi_hal = A4

            for idx, s in enumerate(siswa_terpilih):
                posisi = idx % (KOLOM * BARIS)
                if idx > 0 and posisi == 0:
                    c.showPage()
                kol, bar = posisi % KOLOM, posisi // KOLOM
                x = MARGIN_X + kol * (LEBAR_KARTU + GAP)
                y = tinggi_hal - MARGIN_Y - (bar + 1) * TINGGI_KARTU - bar * GAP

                # Bingkai kartu
                c.setStrokeColor(COLOR_PRIMARY)
                c.setLineWidth(1)
                c.roundRect(x, y, LEBAR_KARTU, TINGGI_KARTU, 4, stroke=1, fill=0)

                # Header
                c.setFillColor(COLOR_PRIMARY)
                c.rect(x, y + TINGGI_KARTU - 0.7 * cm, LEBAR_KARTU, 0.7 * cm, stroke=0, fill=1)
                c.setFillColor(colors.white)
                c.setFont("Helvetica-Bold", 8)
                c.drawCentredString(x + LEBAR_KARTU / 2, y + TINGGI_KARTU - 0.48 * cm,
                                     "KARTU ABSEN SISWA")

                # Kotak foto (kiri)
                kotak_x, kotak_y = x + 0.3 * cm, y + 0.35 * cm
                kotak_w, kotak_h = 2.3 * cm, 3.0 * cm
                c.setDash(2, 2)
                c.setStrokeColor(colors.grey)
                c.rect(kotak_x, kotak_y, kotak_w, kotak_h, stroke=1, fill=0)
                c.setDash()
                c.setFillColor(colors.grey)
                c.setFont("Helvetica", 6)
                c.drawCentredString(kotak_x + kotak_w / 2, kotak_y + kotak_h / 2, "Tempel")
                c.drawCentredString(kotak_x + kotak_w / 2, kotak_y + kotak_h / 2 - 0.25 * cm, "Foto")

                # Info siswa (kanan foto)
                info_x = kotak_x + kotak_w + 0.3 * cm
                info_w = LEBAR_KARTU - (info_x - x) - 0.25 * cm
                c.setFillColor(colors.black)
                c.setFont("Helvetica-Bold", 9)
                c.drawString(info_x, y + TINGGI_KARTU - 1.15 * cm, (s["Nama"][:22]))
                c.setFont("Helvetica", 7.5)
                c.drawString(info_x, y + TINGGI_KARTU - 1.55 * cm, f"Kelas: {a_kelas}")
                c.drawString(info_x, y + TINGGI_KARTU - 1.9 * cm, f"NISN: {s.get('NISN', '')}")
                c.setFont("Helvetica-Oblique", 6.5)
                c.drawString(info_x, y + TINGGI_KARTU - 2.25 * cm, st.session_state.sekolah[:26])

                # QR code, pojok kanan bawah (persegi, lebih keren & lebih mudah discan dari
                # berbagai sudut dibanding barcode batang)
                try:
                    qr_buf = buat_gambar_qr_nisn(str(s.get("NISN", "")).strip())
                    img = ImageReader(qr_buf)
                    qr_sisi = 2.5 * cm
                    qr_x = x + LEBAR_KARTU - qr_sisi - 0.3 * cm
                    qr_y = y + 0.2 * cm
                    c.drawImage(img, qr_x, qr_y, width=qr_sisi, height=qr_sisi,
                                preserveAspectRatio=True, mask="auto")
                except Exception:
                    c.setFont("Helvetica", 6)
                    c.drawString(x + 0.3 * cm, y + 0.3 * cm, "(QR gagal: NISN tidak valid)")

            c.save()
            buf.seek(0)
            st.success(f"✨ {len(siswa_terpilih)} kartu QR berhasil dibuat!")
            st.download_button("⬇️ Unduh PDF Kartu QR", buf.getvalue(),
                                f"Kartu_QR_{a_kelas.replace(' ', '_')}.pdf",
                                "application/pdf", use_container_width=True)

with sub_scan:
    st.subheader("📷 Scan Absen (QR Code NISN)")
    if not (AUTH_AKTIF and st.session_state.user):
        st.warning("⚠️ Fitur ini butuh akun (login) karena kehadiran tersimpan per guru di database. "
                   "Silakan masuk/daftar dulu.")
    else:
        user_id = st.session_state.user["id"]
        access_token = st.session_state.user.get("access_token", "")
        c1, c2 = st.columns(2)
        with c1:
            scan_tanggal = st.date_input("Tanggal Absen", value=date.today(), key="scan_tanggal")
        with c2:
            scan_kelas = st.selectbox("Kelas", KELAS_OPTIONS, key="scan_kelas")
        tanggal_str = scan_tanggal.strftime("%Y-%m-%d")

        peta_nisn = {str(s.get("NISN", "")).strip(): s["Nama"]
                     for s in st.session_state.siswa_rows if str(s.get("NISN", "")).strip()}

        def _proses_scan(nisn_masuk: str):
            nisn_masuk = nisn_masuk.strip()
            if not nisn_masuk:
                return
            nama_siswa = peta_nisn.get(nisn_masuk, f"(NISN {nisn_masuk} — tidak ada di data siswa)")
            ok, err = db_catat_absensi_barcode(user_id, nisn_masuk, nama_siswa, scan_kelas,
                                                 tanggal_str, "Hadir")
            if ok:
                st.session_state["_scan_terakhir"] = f"✅ {nama_siswa} tercatat Hadir."
            else:
                st.session_state["_scan_terakhir"] = f"❌ Gagal mencatat: {err}"

        st.markdown("### 1️⃣ Alat Scanner Fisik (USB/Bluetooth)")
        st.caption("Sambungkan scanner ke HP/laptop, klik kotak di bawah, lalu tembak QR di "
                   "kartu — scanner otomatis 'mengetik' NISN + Enter, siap langsung untuk kartu berikutnya.")
        if "_scan_counter" not in st.session_state:
            st.session_state["_scan_counter"] = 0

        def _on_scan_fisik():
            key = f"scan_fisik_{st.session_state['_scan_counter']}"
            _proses_scan(st.session_state.get(key, ""))
            st.session_state["_scan_counter"] += 1

        st.text_input("Arahkan kursor ke sini lalu scan QR",
                       key=f"scan_fisik_{st.session_state['_scan_counter']}",
                       on_change=_on_scan_fisik, placeholder="Menunggu scan...")
        if st.session_state.get("_scan_terakhir"):
            st.info(st.session_state["_scan_terakhir"])

        st.divider()
        st.markdown("### 2️⃣ Kamera HP")
        st.caption("Izinkan akses kamera saat diminta browser. Arahkan kamera ke QR code di kartu. "
                   "Setiap berhasil scan, kehadiran **langsung tersimpan otomatis** dan kamera tetap "
                   "menyala — tinggal arahkan ke kartu berikutnya, tanpa perlu buka-tutup kamera lagi.")

        import json as _json
        konfigurasi_js = _json.dumps({
            "supabaseUrl": SUPABASE_URL,
            "anonKey": SUPABASE_KEY,
            "accessToken": access_token,
            "userId": user_id,
            "kelas": scan_kelas,
            "tanggal": tanggal_str,
            "petaNisn": peta_nisn,
        })

        html_scanner = f"""
        <style>
          #reader_wrap {{ font-family: sans-serif; }}
          #reader {{ width:100%; max-width:420px; min-height:280px; background:#000; }}
          #reader video {{ width:100% !important; height:auto !important; display:block !important; }}
          #reader img {{ display:none !important; }}
          #btn_mulai_kamera {{
            background:#1E4D8C; color:#fff; border:none; padding:10px 16px;
            border-radius:6px; font-size:14px; cursor:pointer; margin-bottom:10px;
          }}
          #hasil_scan {{ margin-top:8px; font-size:14px; font-weight:bold; }}
          #log_scan {{ margin-top:6px; font-size:12.5px; color:#333; max-height:140px; overflow-y:auto; }}
          #log_scan div {{ padding:2px 0; border-bottom:1px solid #eee; }}
        </style>
        <div id="reader_wrap">
          <button id="btn_mulai_kamera">▶️ Mulai Kamera</button>
          <div id="reader"></div>
          <div id="hasil_scan"></div>
          <div id="log_scan"></div>
        </div>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/html5-qrcode/2.3.8/html5-qrcode.min.js"></script>
        <script>
        const CFG = {konfigurasi_js};
        let terakhirKode = null;
        let terakhirWaktu = 0;
        let totalSesi = 0;

        async function simpanKehadiran(nisn) {{
            const endpoint = CFG.supabaseUrl + "/rest/v1/absensi_barcode?on_conflict=user_id,nisn,tanggal";
            const nama = CFG.petaNisn[nisn] || ("(NISN " + nisn + " — tidak ada di data siswa)");
            const payload = {{
                user_id: CFG.userId, nisn: nisn, nama: nama, kelas: CFG.kelas,
                tanggal: CFG.tanggal, status: "Hadir"
            }};
            try {{
                const resp = await fetch(endpoint, {{
                    method: "POST",
                    headers: {{
                        "Content-Type": "application/json",
                        "apikey": CFG.anonKey,
                        "Authorization": "Bearer " + (CFG.accessToken || CFG.anonKey),
                        "Prefer": "resolution=merge-duplicates,return=minimal"
                    }},
                    body: JSON.stringify(payload)
                }});
                return {{ ok: resp.ok, nama: nama, status: resp.status }};
            }} catch (e) {{
                return {{ ok: false, nama: nama, status: "jaringan" }};
            }}
        }}

        function mulaiScanner() {{
            const el = document.getElementById("hasil_scan");
            const log = document.getElementById("log_scan");
            const btn = document.getElementById("btn_mulai_kamera");
            btn.disabled = true;
            btn.innerText = "🔄 Membuka kamera...";
            const scanner = new Html5Qrcode("reader");
            const config = {{ fps: 10, qrbox: {{ width: 230, height: 230 }} }};
            scanner.start(
                {{ facingMode: "environment" }}, config,
                async (decodedText) => {{
                    const sekarang = Date.now();
                    if (decodedText === terakhirKode && (sekarang - terakhirWaktu) < 4000) {{
                        return; // cegah kartu yang sama ke-scan berkali-kali saat masih di depan kamera
                    }}
                    terakhirKode = decodedText;
                    terakhirWaktu = sekarang;
                    el.innerText = "⏳ Menyimpan " + decodedText + " ...";
                    const hasil = await simpanKehadiran(decodedText);
                    totalSesi += 1;
                    if (hasil.ok) {{
                        el.innerText = "✅ " + hasil.nama + " tercatat Hadir. Siap kartu berikutnya (total sesi ini: " + totalSesi + ")";
                    }} else {{
                        el.innerText = "❌ Gagal menyimpan (" + hasil.status + "). Coba scan ulang.";
                    }}
                    const baris = document.createElement("div");
                    baris.innerText = (hasil.ok ? "✅ " : "❌ ") + hasil.nama + " — " + new Date().toLocaleTimeString();
                    log.prepend(baris);
                }},
                () => {{}}
            ).then(() => {{
                btn.style.display = "none";
            }}).catch((err) => {{
                el.innerText = "❌ Kamera tidak bisa dibuka: " + err;
                btn.disabled = false;
                btn.innerText = "▶️ Coba Lagi";
            }});
        }}
        document.getElementById("btn_mulai_kamera").addEventListener("click", mulaiScanner);
        </script>
        """
        st.components.v1.html(html_scanner, height=560)
        st.caption("Tekan tombol **'Mulai Kamera'** dulu (bukan otomatis) — ini wajib di kebanyakan "
                   "browser HP agar video kamera benar-benar tampil, bukan cuma aktif di belakang layar. "
                   "Kalau tetap tidak muncul, pakai scanner fisik di atas sebagai cadangan.")
        if st.button("🔄 Refresh Daftar Sudah Tercatat", key="btn_refresh_scan"):
            st.rerun()
        st.caption("Daftar di bawah ini diambil dari database, jadi tekan tombol refresh di atas "
                   "kapan pun untuk melihat hasil scan kamera terbaru (kamera menyimpan langsung ke "
                   "database tanpa perlu memuat ulang halaman).")

        st.divider()
        st.markdown("### ✏️ Tandai Manual (Sakit/Izin/Alpa)")
        mc1, mc2, mc3 = st.columns([2, 1, 1])
        with mc1:
            nama_manual = st.selectbox("Siswa", [s["Nama"] for s in st.session_state.siswa_rows]
                                        if st.session_state.siswa_rows else ["(belum ada siswa)"],
                                        key="manual_nama")
        with mc2:
            status_manual = st.selectbox("Status", ["Sakit", "Izin", "Alpa", "Hadir"], key="manual_status")
        with mc3:
            st.write("")
            st.write("")
            if st.button("Catat", key="btn_manual_absen"):
                siswa_manual = next((s for s in st.session_state.siswa_rows if s["Nama"] == nama_manual), None)
                if siswa_manual:
                    nisn_m = str(siswa_manual.get("NISN", "")).strip() or f"NOISN-{siswa_manual.get('id', nama_manual)}"
                    ok, err = db_catat_absensi_barcode(user_id, nisn_m, nama_manual, scan_kelas,
                                                         tanggal_str, status_manual)
                    if ok:
                        st.success(f"✅ {nama_manual} dicatat {status_manual}.")
                        st.rerun()
                    else:
                        st.error(f"❌ Gagal: {err}")

        st.divider()
        st.markdown(f"### 📋 Sudah Tercatat Hari Ini ({scan_tanggal.strftime('%d-%m-%Y')})")
        data_hari_ini = db_ambil_absensi_tanggal(user_id, tanggal_str)
        if data_hari_ini:
            df_hari_ini = pd.DataFrame(data_hari_ini)[["nama", "nisn", "status", "waktu"]]
            df_hari_ini.columns = ["Nama", "NISN", "Status", "Waktu"]
            st.dataframe(df_hari_ini, use_container_width=True, hide_index=True)
            st.caption(f"Total tercatat: {len(data_hari_ini)} siswa")
        else:
            st.info("Belum ada yang tercatat hari ini.")

with sub_rekap:
    st.subheader("📊 Rekap Absen Bulanan")
    if not (AUTH_AKTIF and st.session_state.user):
        st.warning("⚠️ Fitur ini butuh akun (login) karena kehadiran tersimpan per guru di database.")
    else:
        user_id = st.session_state.user["id"]
        rc1, rc2 = st.columns(2)
        with rc1:
            bulan_pilihan = st.selectbox("Bulan", list(range(1, 13)),
                                          index=date.today().month - 1,
                                          format_func=lambda m: [
                                              "Januari", "Februari", "Maret", "April", "Mei", "Juni",
                                              "Juli", "Agustus", "September", "Oktober", "November", "Desember"
                                          ][m - 1], key="rekap_bulan")
        with rc2:
            tahun_pilihan = st.number_input("Tahun", 2020, 2100, date.today().year, key="rekap_tahun")

        tgl_awal = date(tahun_pilihan, bulan_pilihan, 1)
        tgl_akhir = date(tahun_pilihan + (1 if bulan_pilihan == 12 else 0),
                          1 if bulan_pilihan == 12 else bulan_pilihan + 1, 1) - timedelta(days=1)

        if st.button("🔍 Tampilkan Rekap", key="btn_rekap"):
            data_bulan = db_ambil_absensi_rentang(user_id, tgl_awal.strftime("%Y-%m-%d"),
                                                    tgl_akhir.strftime("%Y-%m-%d"))
            if not data_bulan:
                st.info("Belum ada data absensi pada bulan ini.")
            else:
                df = pd.DataFrame(data_bulan)
                hari_efektif = df["tanggal"].nunique()
                pivot = df.pivot_table(index=["nama", "nisn"], columns="status",
                                        values="id", aggfunc="count", fill_value=0).reset_index()
                for kol in ["Hadir", "Sakit", "Izin", "Alpa"]:
                    if kol not in pivot.columns:
                        pivot[kol] = 0
                pivot["Total Tercatat"] = pivot[["Hadir", "Sakit", "Izin", "Alpa"]].sum(axis=1)
                pivot["% Hadir"] = (pivot["Hadir"] / pivot["Total Tercatat"].replace(0, 1) * 100).round(1)
                pivot = pivot.rename(columns={"nama": "Nama", "nisn": "NISN"})
                pivot = pivot[["Nama", "NISN", "Hadir", "Sakit", "Izin", "Alpa", "Total Tercatat", "% Hadir"]]

                st.caption(f"Jumlah hari tercatat dalam bulan ini: **{hari_efektif} hari**")
                st.dataframe(pivot, use_container_width=True, hide_index=True)

                buf_excel = BytesIO()
                with pd.ExcelWriter(buf_excel, engine="openpyxl") as writer:
                    pivot.to_excel(writer, index=False, sheet_name="Rekap Absen")
                st.download_button("⬇️ Unduh Rekap (Excel)", buf_excel.getvalue(),
                                    f"Rekap_Absen_{tahun_pilihan}-{bulan_pilihan:02d}.xlsx",
                                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
